#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
ENM458 Ara Rapor - Tam Eksiksiz DOCX Oluşturucu
Tez Yazım Kılavuzuna uygun: Times New Roman, 12pt, 1.5 satır aralığı,
Sol: 3.5cm, Sağ: 2.5cm, Üst: 3cm, Alt: 2.5cm
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ─── SAYFA AYARLARI ────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Cm(21)
section.page_height = Cm(29.7)
section.top_margin    = Cm(3.0)
section.bottom_margin = Cm(2.5)
section.left_margin   = Cm(3.5)
section.right_margin  = Cm(2.5)

# ─── NORMAL STİL ────────────────────────────────────────────────────────────
normal = doc.styles['Normal']
normal.font.name = 'Times New Roman'
normal.font.size = Pt(12)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after  = Pt(6)
normal.paragraph_format.line_spacing = Pt(18)   # 1.5 satır ≈ 18pt
normal.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.JUSTIFY

# Heading stilleri
for i in range(1, 4):
    h = doc.styles[f'Heading {i}']
    h.font.name  = 'Times New Roman'
    h.font.color.rgb = RGBColor(0, 0, 0)
    h.font.bold  = True
    h.font.size  = Pt(12)
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after  = Pt(3)
    h.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.LEFT

# ─── YARDIMCI FONKSİYONLAR ─────────────────────────────────────────────────

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name  = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.bold  = True
        run.font.size  = Pt(12)
    return h

def add_para(text='', indent=True, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.paragraph_format.alignment    = align
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    if indent:
        p.paragraph_format.first_line_indent = Cm(1)
    if text:
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold   = bold
        run.italic = italic
    return p

def add_mixed(parts, indent=True):
    """parts: list of (text, bold, italic)"""
    p = doc.add_paragraph()
    p.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    if indent:
        p.paragraph_format.first_line_indent = Cm(1)
    for text, bold, italic in parts:
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.bold   = bold
        run.italic = italic
    return p

def add_page_break():
    doc.add_page_break()

def set_cell_style(cell, text, bold=False, size=10, bg=None):
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    if bg:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), bg)
        tcPr.append(shd)

def add_table_caption(text):
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.first_line_indent = Cm(0)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run.bold = True

# ═══════════════════════════════════════════════════════════════════════════
# KAPAK (İÇ KAPAK)
# ═══════════════════════════════════════════════════════════════════════════

p = doc.add_paragraph()
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(0)
run = p.add_run('TİKTOK\'TA ETKİLİ MARKA İŞ BİRLİĞİ İÇİN')
run.font.name = 'Times New Roman'; run.font.size = Pt(12); run.bold = True

p = doc.add_paragraph()
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('İÇERİK ÜRETİCİSİ SEÇİMİNİN OPTİMİZASYONU')
run.font.name = 'Times New Roman'; run.font.size = Pt(12); run.bold = True

p = doc.add_paragraph()
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Lisans Tezi')
run.font.name = 'Times New Roman'; run.font.size = Pt(12)

doc.add_paragraph()

p = doc.add_paragraph()
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Esra ZEYREK\nAbdullah Raif YILDIRIM\nOğuzhan DİKMEN')
run.font.name = 'Times New Roman'; run.font.size = Pt(12)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('ENM458 ENDÜSTRİ MÜHENDİSLİĞİ BİTİRME PROJESİ II')
run.font.name = 'Times New Roman'; run.font.size = Pt(12); run.bold = True

p = doc.add_paragraph()
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('ARA SINAV RAPORU')
run.font.name = 'Times New Roman'; run.font.size = Pt(12); run.bold = True

doc.add_paragraph()

p = doc.add_paragraph()
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Danışman: Asst. Prof. Dr. Zeliha ERGÜL AYDIN')
run.font.name = 'Times New Roman'; run.font.size = Pt(12)

p = doc.add_paragraph()
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Eskişehir Teknik Üniversitesi Mühendislik Fakültesi\nEndüstri Mühendisliği Bölümü\nNisan, 2026')
run.font.name = 'Times New Roman'; run.font.size = Pt(12)

add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# ÖZET
# ═══════════════════════════════════════════════════════════════════════════

add_heading('ÖZET', level=1)

add_para('TİKTOK\'TA ETKİLİ MARKA İŞ BİRLİĞİ İÇİN İÇERİK ÜRETİCİSİ SEÇİMİNİN OPTİMİZASYONU', indent=False, bold=True)
add_para('Esra ZEYREK | Abdullah Raif YILDIRIM | Oğuzhan DİKMEN', indent=False)
add_para('Endüstri Mühendisliği — Eskişehir Teknik Üniversitesi, Mühendislik Fakültesi, Nisan, 2026\nDanışman: Asst. Prof. Dr. Zeliha ERGÜL AYDIN', indent=False)

add_para('Dijital pazarlama ekosisteminde kısa video temelli içeriklerin önem kazanması, özellikle TikTok platformunu marka iletişim stratejilerinin merkezine taşımıştır. Ancak influencer seçimi çoğu zaman takipçi sayısı gibi sınırlı göstergelere dayandırılmakta, bu durum reklam bütçelerinin etkin kullanılamamasına ve yanlış hedef kitle eşleşmelerine neden olmaktadır.')

add_para('Bu çalışma, TikTok üzerinde gerçekleştirilecek marka iş birlikleri için veri temelli ve matematiksel bir içerik üreticisi seçim modeli geliştirmeyi amaçlamaktadır. Birinci dönemde (ENM457) Selenium tabanlı web kazıma yöntemiyle içerik üreticilerine ait profil ve video metrikleri toplanmış; bu veriler ön işleme ve keşifsel analizden geçirilmiştir. İkinci dönemde (ENM458) ise toplanan veriler üzerinden etkileşim vekilleri (engagement proxy) hesaplanmış, takipçi bazlı kademeli bir maliyet tahmin modeli kurulmuş ve çok kriterli karar verme (MCDM) skoru üretilmiştir. Elde edilen bu sayısal göstergeler, Gurobi çözücüsü kullanılarak kurulan 0-1 Tam Sayılı Doğrusal Programlama (ILP) modeline girdi olarak aktarılmıştır. Model; bütçe kısıtı, minimum çeşitlilik, kategori dengesi ve kara liste toleransı gibi gerçekçi iş koşullarını barındırmakta ve Streamlit tabanlı interaktif bir karar destek arayüzü ile son kullanıcıya sunulmaktadır.')

add_para('Elde edilen bulgular, modelin farklı bütçe senaryolarında tutarlı ve anlamlı influencer portföyleri önerdiğini, takipçi büyüklüğü yerine etkileşim kalitesini ön plana alan çözümler ürettiğini ortaya koymaktadır. Bu yaklaşım, dijital reklam kaynaklarının daha sorumlu ve verimli kullanılmasına katkı sağlamakta; SKA 12: Sorumlu Üretim ve Tüketim ile SKA 8: İnsana Yakışır İş ve Ekonomik Büyüme hedefleriyle doğrudan ilişkilendirilmektedir.')

add_mixed([('Anahtar Sözcükler: ', True, False), ('TikTok, içerik üreticisi seçimi, influencer optimizasyonu, tam sayılı doğrusal programlama, MCDM, Gurobi, web kazıma, dijital pazarlama, karar destek sistemi, sorumlu tüketim', False, False)], indent=False)

add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# ABSTRACT
# ═══════════════════════════════════════════════════════════════════════════

add_heading('ABSTRACT', level=1)
add_para('OPTIMIZATION OF CONTENT CREATOR SELECTION FOR EFFECTIVE BRAND COLLABORATIONS ON TIKTOK', indent=False, bold=True)
add_para('Esra ZEYREK | Abdullah Raif YILDIRIM | Oğuzhan DİKMEN', indent=False)
add_para('Department of Industrial Engineering — Eskisehir Technical University, Engineering Faculty, April, 2026\nSupervisor: Asst. Prof. Dr. Zeliha ERGÜL AYDIN', indent=False)

add_para('The growing dominance of short-form video content has positioned TikTok as a central platform in brand communication strategies. However, influencer selection is often based on simplistic metrics such as follower count, leading to inefficient allocation of advertising budgets and poor audience alignment.')
add_para('This study aims to develop a data-driven, mathematical content creator selection model for brand collaborations on TikTok. In the first phase (ENM457), profile and video-level engagement data were collected via Selenium-based web scraping and subjected to exploratory analysis. In the second phase (ENM458), engagement proxies were computed, a tiered cost estimation model was established, and a Multi-Criteria Decision Making (MCDM) score was generated for each creator. These indicators were fed into a 0-1 Integer Linear Programming (ILP) model solved with Gurobi, incorporating realistic business constraints. The system is presented through an interactive Streamlit-based decision support interface.')
add_para('Results indicate that the model consistently generates meaningful influencer portfolio recommendations across different budget scenarios, prioritizing engagement quality over follower volume, directly aligning with SDG 12: Responsible Consumption and Production and SDG 8: Decent Work and Economic Growth.')

add_mixed([('Keywords: ', True, False), ('TikTok, content creator selection, influencer optimization, integer linear programming, MCDM, Gurobi, web scraping, digital marketing, decision support system, responsible consumption', False, False)], indent=False)

add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 1. GİRİŞ
# ═══════════════════════════════════════════════════════════════════════════

add_heading('1. GİRİŞ', level=1)

add_para('Sosyal medya platformlarının küresel kullanıcı kitlesini hızla büyütmesi ve bu platformların gündelik yaşamdaki yerini güçlendirmesiyle birlikte, dijital pazarlama faaliyetleri de köklü bir dönüşüm sürecine girmiştir. Geleneksel reklam mecralarının yerini giderek daha fazla sosyal medya iletişimi alırken, içerik üreticileri (influencer\'lar) bu yeni reklamcılık düzeninin merkezinde konumlanmaktadır. TikTok özelinde bakıldığında, kısa video formatının hızlı tüketim yapısı ve algoritmik içerik keşif mekanizması, platformu markalar için stratejik bir iletişim kanalına dönüştürmüştür.')
add_para('Influencer pazarlamasının büyümesiyle birlikte, içerik üreticisi seçim süreci de giderek daha kritik bir karar noktasına evrilmiştir. Küresel ölçekte influencer pazarlaması hacminin 2023 yılı itibarıyla 21 milyar doları aştığı tahmin edilmektedir [3]. Bu büyüme, aynı zamanda bütçe verimsizliği, yanlış hedef kitle eşleşmesi ve ölçümleme güçlükleri gibi sorunları beraberinde getirmektedir. Özellikle Türkiye ölçeğinde influencer seçimi çoğunlukla sezgisel değerlendirmelere, ajans ağlarına veya yalnızca takipçi sayısı gibi yüzeysel kriterlere dayandırılmaktadır.')
add_para('Bu gerçeklik, projenin temel motivasyonunu oluşturmaktadır: influencer seçimini "veri temelli bir optimizasyon problemi" olarak yeniden tanımlamak. Çalışmanın birinci döneminde (ENM457), bu problemi çözmek için gereken veri altyapısı Selenium tabanlı web kazıma yöntemiyle oluşturulmuştur. TikTok\'un Türkiye\'de resmi API erişimine kapalı olması nedeniyle veriler herkese açık profil sayfalarından sistematik biçimde toplanmış; profil bilgileri, video bazlı etkileşim metrikleri ve içerik tanımlayıcıları (hashtag, biyografi, kategori) bir araya getirilmiştir.')
add_para('Çalışmanın ikinci dönemi olan ENM458 kapsamında ise bu ham veriler işlenmiş, analitik bileşenlere dönüştürülmüş ve nihayet bir optimizasyon modeline entegre edilmiştir. Gurobi akademik lisansıyla çözülen 0-1 Tam Sayılı Doğrusal Programlama (ILP) modeli, markanın bütçe kısıtları ve stratejik tercihleri doğrultusunda en yüksek performanslı influencer portföyünü otomatik olarak belirlemektedir. Sistem, Streamlit çerçevesiyle geliştirilen interaktif bir karar destek arayüzüyle son kullanıcıya sunulmaktadır.')
add_para('Bu rapor, ENM458 Ara Sınav Raporunu (1-6. haftalar) kapsamakta olup modelin dayandığı yöntemi, veri hazırlık sürecini ve gerçekleştirilen uygulamayı bütünleşik biçimde aktarmaktadır.')

# 1.1
add_heading('1.1. Literatür Taraması', level=2)

add_para('Influencer pazarlaması, hem pazarlama hem de bilgisayar bilimi literatüründe giderek artan bir ilgi görmektedir. Johnson ve Alvarez [1], dijital içerik üreticilerinin markalar ile tüketiciler arasında güven temelli bir köprü kurduğunu ve geleneksel reklamcılıktan farklı bir etkileşim yapısı sunduğunu belirtmektedir. Kumar ve Lee [2] ise içerik üreticisi pazarlamasının kullanıcı etkileşimi ve içerik özgünlüğü üzerine kurulu yapısını incelemiştir.')
add_para('Zhang ve Han [3], influencer pazarlamasının teorik temellerini, güncel eğilimlerini ve gelecekteki araştırma yönlerini kapsamlı biçimde ele almıştır. Sektörün 2023 itibarıyla 21 milyar dolara ulaştığını vurgulayan bu çalışma, alanın akademik açıdan olgunlaşmakta olduğuna işaret etmektedir. Andersson ve Bergström [4] ise influencer başarısının yalnızca takipçi sayısıyla ölçülemeyeceğini; içerik kalitesi ve etkileşim yapısının da belirleyici olduğunu ampirik bulgularla ortaya koymuştur.')
add_para('Bütçe kısıtlı optimizasyon yaklaşımları açısından Kim ve Cho [10], influencer kampanyalarının matematiksel modelleme ile optimize edilebileceğini göstermiştir. López-Dawn ve Giovanidis [11] ise influencer seçimini bütçe kısıtlı bir portföy optimizasyon problemi olarak ele alarak sistematik bir çözüme kavuşturmuştur. Ancak bu çalışmaların büyük çoğunluğu Instagram veya YouTube gibi platformlara odaklanmakta; TikTok\'un kendine özgü algoritmik yapısını ve Türkiye koşullarını yansıtmamaktadır.')
add_para('De Veirman ve arkadaşları [5], yüksek takipçi sayısının her zaman yüksek reklam etkisi anlamına gelmediğini göstermiştir. Bu bulgu, projemizin takipçi sayısı yerine etkileşim kalitesini merkeze alan tasarımını doğrudan desteklemektedir. Lou ve Yuan [6], influencer mesajlarının algılanan değeri ve güvenilirliğinin tüketici güvenini doğrudan etkilediğini belirtmiştir. Saito ve Kobayashi [7], influencer maksimizasyon problemini matematiksel bir çerçeveye taşımış; Li ve Wang [9] ise büyük veri analitiği ve ağ tabanlı yöntemlerin influencer etkinliğini anlamakta önemli katkılar sunduğunu ortaya koymuştur.')
add_para('Tüm bu literatür tarandığında, TikTok verisine dayalı, Türkiye koşullarına özgü ve gerçekçi iş kısıtlarını içeren bir optimizasyon modelinin henüz geliştirilmediği görülmektedir. Bu çalışma, tam da bu boşluğu hedef alarak özgün bir katkı sunmayı amaçlamaktadır.')

# 1.2
add_heading('1.2. Problemin Kısaca Tarifi', level=2)
add_para('Projenin çözdüğü problem şu şekilde özetlenebilir: Bir marka, belirli bir bütçe dahilinde TikTok\'ta reklam kampanyası yürütmek istemektedir. Bütçesini en verimli şekilde kullanarak maksimum reklam etkisi yaratacak influencer kombinasyonunu nasıl seçmelidir?')
add_para('Bu problem birden fazla alt bileşeni içermektedir. Veri sorunu açısından, TikTok Türkiye\'de resmi API erişimi sunmamaktadır. Bu nedenle influencer verileri özgün web kazıma araçlarıyla toplanmak zorundadır. Ölçüm sorunu açısından, bir influencer\'ın "ne kadar iyi" olduğunu tek bir sayıyla ifade etmek mümkün değildir; bunu çözmek için çok kriterli bir değerlendirme skoru (MCDM skoru) geliştirilmiştir. Maliyet belirsizliği açısından, influencer fiyatları kamuya açık değildir; sektör benchmarkları ve takipçi kademeleri kullanılarak gerçekçi bir maliyet tahmin modeli oluşturulmuştur. Karar sorunu açısından, yüzlerce aday arasından bütçeye uygun, kategorileri dengeli, etkileşim kalitesi yüksek bir portföy oluşturmak NP-zor bir kombinatoryal optimizasyon problemidir; bu sorun 0-1 ILP modeliyle çözülmüştür. Son olarak kullanılabilirlik sorunu açısından, modelin uygulama değeri yaratması için teknik olmayan son kullanıcıların da erişebileceği bir arayüze ihtiyaç vardır; bu arayüz Streamlit ile geliştirilmiştir.')
add_para('Problemin gerçekçi kısıtları incelendiğinde ekonomik, stratejik, itibar yönetimi ve risk yönetimi boyutları öne çıkmaktadır. Ekonomik açıdan reklam bütçesi sınırlıdır ve her influencer\'ın maliyeti farklıdır. Stratejik açıdan marka, belirli kategorilerde mutlaka temsilci istemektedir. İtibar yönetimi açısından marka, bazı influencer\'larla çalışmak istememektedir (kara liste). Risk yönetimi açısından ise tüm bütçenin tek bir hesaba verilmesi hem finansal hem de PR riski taşımaktadır (çeşitlilik kısıtı).')

# 1.3
add_heading('1.3. Projenin Girişimcilik ve Yenilikçilik Açısından Katkısı', level=2)
add_para('Bu proje, salt akademik bir çalışmanın sınırlarını aşarak ticarileşme potansiyeli yüksek bir girişim prototipi sunmaktadır. Türkiye\'deki dijital pazarlama ajanslarının büyük çoğunluğu, influencer seçimini spreadsheet\'ler, kişisel bağlantılar ve sezgilerle yapmaktadır. Bu proje, söz konusu boşluğu dolduracak, matematiksel altyapıya dayanan, ölçeklenebilir ve yeniden kullanılabilir bir karar destek aracı sunmaktadır.')
add_para('Geliştirilen Streamlit arayüzü, aylık abonelik modeliyle (B2B SaaS) dijital ajanslar ve markalara sunulabilecek bir minimum uygulanabilir ürün (MVP) niteliği taşımaktadır. Kullanıcı, bütçesini ve tercihlerini girerek saniyeler içinde optimize edilmiş bir influencer listesi alabilmektedir. Yenilikçi unsurlar arasında Türkiye\'de TikTok verisine dayalı ilk optimizasyon modeli olma özelliği, API erişimi gerektirmeyen web kazıma tabanlı özgün veri toplama altyapısı, takipçi sayısı yerine etkileşim kalitesini ön plana alan MCDM skorlama sistemi ve gerçek zamanlı senaryo simülasyonu sayılabilir. Proje, Endüstri Mühendisliği, Bilgisayar Mühendisliği ve İşletme disiplinlerini başarıyla entegre eden multidisipliner bir yapıya sahiptir.')

# 1.4
add_heading('1.4. Projenin BM Sürdürülebilir Kalkınma Amaçları Kapsamında Etkileri', level=2)
add_para('SKA 8 — İnsana Yakışır İş ve Ekonomik Büyüme kapsamında değerlendirildiğinde, dijital reklam dünyasındaki mevcut yapıda reklam harcamalarının büyük bölümü "Mega" düzeydeki (1 milyondan fazla takipçili) içerik üreticilerinde yoğunlaşmaktadır. Bu durum, organik etkileşim oranları yüksek ancak henüz tanınmamış "Nano" (0-10 bin takipçi) ve "Micro" (10 bin - 100 bin takipçi) düzeydeki yaratıcıların ekosistemin dışında kalmasına neden olmaktadır. Geliştirilen optimizasyon modeli, maliyet etkinliği kriterini sisteme dahil ederek bu dengesizliği düzeltmektedir. Yüksek etkileşim kalitesine sahip küçük ölçekli içerik üreticileri çok boyutlu bir değerlendirmeden geçirildiğinde algoritma tarafından öne çıkarılabilmekte; KOBİ\'ler ve yerel markalar da uygun maliyetli ama etkili işbirliği yapma imkânı bulmaktadır.')
add_para('SKA 12 — Sorumlu Üretim ve Tüketim kapsamında değerlendirildiğinde, dijital reklam bütçelerinin sezgisel kararlar veya tekil metriklerle tahsis edilmesi kaynakların israfına ve başarısız kampanyaların tekrarlanmasına yol açmaktadır. Bu proje, her harcanan reklam lirasının matematiksel hesap görmeyi zorunlu kılan bir sistem kurmaktadır. MCDM skoru, bütçe kullanım oranı ve seçim gerekçeleri şeffaf ve sayısal biçimde sunulmakta; sorumlu tüketim teşvik edilmekte ve dijital reklam sektöründe sürdürülebilir bir karar verme kültürünün yerleşmesine katkı sağlanmaktadır.')

add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 2. YÖNTEM
# ═══════════════════════════════════════════════════════════════════════════

add_heading('2. TAM SAYILI DOĞRUSAL PROGRAMLAMA YÖNTEMİ', level=1)

add_para('Tam Sayılı Doğrusal Programlama (Integer Linear Programming — ILP), karar değişkenlerinin tamsayı (tam sayı veya ikili 0-1) değerler almasını zorunlu kılan doğrusal optimizasyon problemlerinin çözümüne yönelik geliştirilmiş bir yöntemdir. ILP modelleri, hem amaç fonksiyonu hem de kısıtlar bakımından doğrusal ilişkiler içerdiğinden sistematik çözüm algoritmalarına elverişlidir [10, 11].')
add_para('Influencer seçim problemi, doğası gereği ikili bir karar yapısına sahiptir: bir influencer ya seçilir (xᵢ = 1) ya da seçilmez (xᵢ = 0). Bu yapı, problemi doğrudan 0-1 ILP\'ye dönüştürmektedir. Sürekli çözüm yöntemlerine (LP relaxation) kıyasla ILP, gerçekçi ve uygulanabilir portföyler üretmesi bakımından üstündür. Kim ve Cho [10] ile López-Dawn ve Giovanidis [11], benzer yapıdaki influencer seçim problemlerini matematiksel programlama yöntemiyle çözmenin kampanya etkinliğini anlamlı biçimde artırdığını göstermiştir.')
add_para('Gurobi Solver, endüstriyel ölçekli optimizasyon problemleri için geliştirilmiş, akademik ve ticari kullanımda yaygın bir matematiksel programlama çözücüsüdür. Dal-ve-sınır (Branch-and-Bound) algoritması temelinde çalışan Gurobi, LP gevşemesi aracılığıyla tamsayılı çözümler için üst sınır hesaplamakta; ardından kural tabanlı dallara ayırma stratejileriyle arama alanını daraltmaktadır. Çoklu iş parçacığı (Threads=4) ve agresif ön-çözme (Presolve=2) parametreleriyle yapılandırılan model, birkaç yüz değişken içeren influencer problemlerini genellikle saniyeler içinde optimal ya da yakın-optimal biçimde çözmektedir.')
add_para('ILP modeli, belirlenen kısıtlar ve hedef fonksiyonu çerçevesinde matematiksel olarak optimal çözüm sunmaktadır. Ancak bu optimallik, modele yansıtılan varsayımlara bağımlıdır. Maliyet tahmin modeli gerçek fiyatların bir yaklaşımıdır; MCDM skoru geçmiş veri üzerinden türetilmiş olup gelecek performansı kesin olarak öngöremez. Bu sınırlılıklar, modelin bir ön eleme ve karar destek aracı olarak konumlandırılmasını ve uzman insan değerlendirmesiyle desteklenmesini gerektirmektedir.')

add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 3. UYGULAMA
# ═══════════════════════════════════════════════════════════════════════════

add_heading('3. UYGULANMASI: TİKTOK VERİSİYLE ILP MODELİNİN GELİŞTİRİLMESİ', level=1)

add_heading('3.1. Sistem ve Problem Hakkında Genel Bilgiler', level=2)
add_para('Çalışma, TikTok platformunda marka iş birliği yapacak influencer\'ların seçimini konu almaktadır. Türkiye\'de 20 milyonu aşkın aktif TikTok kullanıcısı bulunmakla birlikte, platform Türkiye için resmi API erişimi sunmamaktadır. Bu durum, veri toplama sürecini doğrudan etkilemekte ve özgün mühendislik çözümleri gerektirmektedir.')
add_para('Birinci dönemde (ENM457) yürütülen web kazıma süreci aracılığıyla Türkiye\'deki çeşitli kategorilerde faaliyet gösteren içerik üreticilerine ait profil ve video verileri toplanmıştır. Elde edilen veri seti; kullanıcı adı, takipçi sayısı, toplam beğeni, video sayısı ve video bazlı etkileşim metrikleri ile kategori bilgilerini kapsamaktadır. İçerik üreticileri; Güzellik ve Bakım, Moda ve Stil, Teknoloji ve Dijital, Yiyecek ve Yemek, Komedi ve Eğlence, Spor ve Sağlık, Seyahat ve Yaşam, Eğitim ve Bilgi, Müzik ve Performans ile Oyun kategorileri altında sınıflandırılmıştır.')

add_heading('3.2. Veri Setinin İkinci Döneme Hazırlanması: Analitik Dönüşüm Adımları', level=2)
add_para('Ham veri doğrudan optimizasyon modeline aktarılamaz. Bu nedenle birinci dönemde toplanan veriler, ikinci dönemde bir dizi analitik dönüşüm adımından geçirilmiştir.')

add_heading('3.2.1. Etkileşim Vekili (Engagement Proxy) Hesaplanması', level=3)
add_para('Yalnızca takipçi sayısına dayalı değerlendirmeler yanıltıcı olabilmektedir. Büyük kitleye ulaşabilen ancak düşük etkileşim oranına sahip hesaplar, marka iletişimi açısından sınırlı değer üretmektedir. Bu gerçekten yola çıkılarak, her içerik üreticisi için iki bileşenli bir Etkileşim Vekili skoru türetilmiştir.')
add_para('Bileşen 1 — Takipçi Başına Beğeni Oranı: Her içerik üreticisinin video başına ortalama beğeni sayısı, takipçi sayısına bölünerek hesaplanmıştır (like_per_follower = avg_video_likes / followers). Bileşen 2 — Ortalama Video Beğeni Logaritması: Geniş ölçekli değer farklılıklarını düzenlemek için ortalama video beğenisinin doğal logaritması alınmıştır (log_avg_likes = ln(1 + avg_video_likes)).')
add_para('Etkileşim Vekili (Normalize Edilmiş): Engagement Proxy = 0.60 × MinMax(like_per_follower) + 0.40 × MinMax(log_avg_likes). %60 ağırlık, organik sadık kitleye sahip küçük üreticilerin adaletli biçimde temsil edilmesini sağlarken; %40 ağırlık, mutlak büyüklük avantajını da değerlendirmeye katmaktadır.')

add_heading('3.2.2. Tahmini Maliyet Modeli (Cost Estimation Model)', level=3)
add_para('TikTok işbirliği ücretleri kamuya açık değildir. Sektör benchmarkları ve akademik kaynaklara dayanan kademeli (tier-based) bir maliyet modeli geliştirilmiştir. Temel maliyet formülü şu şekilde işletilmektedir: Maliyet = (Taban Ücret + Takipçi/1000 × Artış Oranı) × Kategori Çarpanı.')

add_table_caption('Tablo 3.1. Takipçi Bazlı Kademeli Maliyet Modeli')
t = doc.add_table(rows=6, cols=4)
t.style = 'Table Grid'
t.alignment = WD_TABLE_ALIGNMENT.CENTER
hdrs = ['Kademe (Tier)', 'Takipçi Aralığı', 'Taban Ücret (TL)', 'Artış Oranı (TL / 1K)']
for i, h in enumerate(hdrs):
    set_cell_style(t.rows[0].cells[i], h, bold=True, size=10, bg='D9D9D9')
data = [
    ('Nano',  '0 – 10.000',          '1.500',  '150'),
    ('Micro', '10.000 – 100.000',    '8.000',  '80'),
    ('Mid',   '100.000 – 500.000',   '35.000', '35'),
    ('Macro', '500.000 – 1.000.000', '90.000', '20'),
    ('Mega',  '1.000.000+',          '200.000','8'),
]
for r, row in enumerate(data):
    for c, val in enumerate(row):
        set_cell_style(t.rows[r+1].cells[c], val, size=10)

doc.add_paragraph()

add_para('Kategori maliyet çarpanları, içerik üreticilerinin faaliyet gösterdiği sektöre göre taban maliyetin üzerine eklenmektedir. Bu yapı, farklı kategorilerdeki arz-talep dengesini ve markaların ödeme eğilimlerini modele yansıtmaktadır.')

add_table_caption('Tablo 3.2. Kategori Maliyet Çarpanları')
t2 = doc.add_table(rows=12, cols=2)
t2.style = 'Table Grid'
t2.alignment = WD_TABLE_ALIGNMENT.CENTER
set_cell_style(t2.rows[0].cells[0], 'Kategori', bold=True, size=10, bg='D9D9D9')
set_cell_style(t2.rows[0].cells[1], 'Maliyet Çarpanı', bold=True, size=10, bg='D9D9D9')
cats = [
    ('Beauty & Personal Care (Güzellik ve Bakım)', '1.15'),
    ('Fashion & Style (Moda ve Stil)', '1.10'),
    ('Technology & Digital (Teknoloji)', '1.05'),
    ('Travel & Lifestyle (Seyahat)', '1.05'),
    ('Fitness & Health (Spor ve Sağlık)', '1.00'),
    ('Music & Performance (Müzik)', '1.00'),
    ('Mixed/Unclear', '1.00'),
    ('Food & Cooking (Yemek)', '0.95'),
    ('Comedy & Entertainment (Komedi)', '0.90'),
    ('Gaming (Oyun)', '0.90'),
    ('Education & Informative (Eğitim)', '0.85'),
]
for r, (cat, mul) in enumerate(cats):
    set_cell_style(t2.rows[r+1].cells[0], cat, size=10)
    set_cell_style(t2.rows[r+1].cells[1], mul, size=10)

doc.add_paragraph()

add_heading('3.2.3. Çok Kriterli Karar Verme (MCDM) Skoru', level=3)
add_para('Birden fazla kriterin bütünleşik biçimde değerlendirilmesi için ağırlıklı doğrusal birleştirme yöntemi kullanılmıştır. Her kriter önce Min-Max normalizasyonuyla [0, 1] aralığına çekilmiş, ardından belirlenen ağırlıklarla ağırlıklı toplamı alınmıştır.')

add_table_caption('Tablo 3.3. MCDM Kriter Ağırlık Tablosu')
t3 = doc.add_table(rows=6, cols=3)
t3.style = 'Table Grid'
t3.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Kriter', 'Açıklama', 'Ağırlık']):
    set_cell_style(t3.rows[0].cells[i], h, bold=True, size=10, bg='D9D9D9')
kriter = [
    ('Etkileşim Skoru',        'Engagement Proxy',              '%35'),
    ('Takipçi Büyüklüğü',      'log(1 + followers)',            '%20'),
    ('Takipçi Başına Beğeni',  'like_per_follower',             '%20'),
    ('Maliyet Etkinliği',      'log(avg_likes) / tahmini maliyet', '%15'),
    ('Video Hacmi',            'Video üretim sayısı',           '%10'),
]
for r, row in enumerate(kriter):
    for c, val in enumerate(row):
        set_cell_style(t3.rows[r+1].cells[c], val, size=10)

doc.add_paragraph()
add_para('MCDM skoru hesaplama formülü: MCDM_Score = 0.35 × c1 + 0.20 × c2 + 0.20 × c3 + 0.15 × c4 + 0.10 × c5. Bu ağırlık yapısının temel tasarım mantığı şudur: etkileşim kalitesi (%55 birleşik ağırlıkla) en güçlü belirleyici unsur iken takipçi sayısının salt büyüklüğü tek başına yeterli kabul edilmemektedir. Maliyet etkinliği (%15) bütçe kullanımını optimize ederken video hacmi (%10) aktif içerik üretimini ödüllendirmektedir.')

add_heading('3.3. ILP Modelinin Uygulanması ve Gurobi Entegrasyonu', level=2)

add_heading('3.3.1. Matematiksel Model Formülasyonu', level=3)
add_para('MCDM skorları ve maliyet tahminleri hesaplandıktan sonra, influencer seçim problemi bir 0-1 Tam Sayılı Doğrusal Programlama (ILP) modeli olarak kurulmuştur.')
add_para('Karar Değişkeni: xᵢ ∈ {0, 1}, ∀i ∈ {1, 2, …, n}. xᵢ = 1 ise i. içerik üreticisi portföye dahil edilir; xᵢ = 0 ise seçilmez.')
add_para('Amaç Fonksiyonu (Maksimizasyon): Z = Σ MCDM_Scoreᵢ × xᵢ. Seçilen influencer\'ların kümülatif MCDM skoru maksimize edilmektedir.')
add_para('Kısıt 1 — Bütçe Kısıtı: Σ costᵢ × xᵢ ≤ B. Seçilen influencer\'ların toplam tahmini maliyeti, önceden belirlenen kampanya bütçesi B\'yi aşamaz.')
add_para('Kısıt 2 — Minimum Çeşitlilik: Σ xᵢ ≥ n_min. Risk dağıtımı ve kampanya çeşitliliği için en az n_min influencer seçimi zorunludur. Tüm bütçenin tek bir hesaba yatırılması hem finansal hem de PR riski taşımaktadır.')
add_para('Kısıt 3 — Maksimum Portföy Büyüklüğü (İsteğe Bağlı): Σ xᵢ ≤ n_max. Yönetilebilirlik açısından maksimum influencer sayısı sınırlandırılabilir.')
add_para('Kısıt 4 — Tier Bazlı Üst Sınır (İsteğe Bağlı): Belirli bir kademedeki influencer sayısı sınırlandırılarak portföy risk dengesi yönetilir.')
add_para('Kısıt 5 — Zorunlu Kategori Temsili (İsteğe Bağlı): Kullanıcı tarafından belirlenen sektör kategorilerinden en az bir influencer seçilmesi sağlanır.')
add_para('Kısıt 6 — Kara Liste (Blacklist): xₖ = 0, ∀k ∈ K_black. Kara listede yer alan influencer\'lar hiçbir koşulda seçilmez. Bu kısıt, markanın itibar yönetimi kararlarını matematiksel garanti altına almaktadır.')

add_heading('3.3.2. Gurobi Parametre Yapılandırması', level=3)

add_table_caption('Tablo 3.4. Gurobi Çözücü Parametre Tablosu')
t4 = doc.add_table(rows=6, cols=3)
t4.style = 'Table Grid'
t4.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Parametre', 'Değer', 'Açıklama']):
    set_cell_style(t4.rows[0].cells[i], h, bold=True, size=10, bg='D9D9D9')
params = [
    ('TimeLimit', '60 saniye', 'Maksimum çözüm süresi'),
    ('MIPGap',    '0.001',     '%0.1 optimallik toleransı'),
    ('Threads',   '4',         'Paralel iş parçacığı sayısı'),
    ('MIPFocus',  '1',         'İyi çözüm bulmaya odaklan'),
    ('Presolve',  '2',         'Agresif ön-çözme'),
]
for r, row in enumerate(params):
    for c, val in enumerate(row):
        set_cell_style(t4.rows[r+1].cells[c], val, size=10)

doc.add_paragraph()
add_para('Model, optimal çözüme ulaştığında (Optimal), zaman limitine yakın çözüm bulduğunda (TimeLimit) veya uygulanabilir olmadığını tespit ettiğinde (Infeasible) olmak üzere farklı sonuç durumlarını raporlamakta ve arayüze aktarmaktadır.')

add_heading('3.4. Streamlit Tabanlı Karar Destek Arayüzü', level=2)
add_para('Gurobi modeli, Tiktokoptimizerphase2v3gurobi.py modülü içinde kapsüllenmiştir. Bu modül, streamlit_app.py aracılığıyla interaktif bir web uygulamasına entegre edilmektedir.')
add_para('Arayüz özellikleri şu şekilde sıralanabilir. Bütçe Slaydırı aracılığıyla kullanıcı, kampanya bütçesini 1.000 TL ile 10.000.000 TL arasında dinamik olarak belirleyebilir. Kategori Seçimi ile zorunlu kategori kısıtları çoklu seçim kutucuklarıyla tanımlanabilir. Kara Liste Yönetimi sayesinde istenmediği düşünülen influencer\'lar listeden seçilerek otomatik olarak model dışında tutulur; kara listedeki hesapların profil görselleri de arayüzde görüntülenmektedir. Sonuç Görüntüleme ekranında optimizasyon çalıştırıldığında; seçilen kişi sayısı, harcanan bütçe, kalan bütçe ve toplam MCDM skoru metrik kartlarda özetlenmektedir. Seçilen influencer\'lar ızgara düzeninde fotoğrafları, kategorileri, takipçi sayıları, tahmini maliyetleri ve MCDM skorlarıyla birlikte listelenmektedir.')

add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# 4. SONUÇ
# ═══════════════════════════════════════════════════════════════════════════

add_heading('4. SONUÇ VE ÖNERİLER', level=1)

add_para('Bu çalışmanın ikinci dönem (ENM458) ilk altı haftasında gerçekleştirilen çalışmalar ve ulaşılan sonuçlar şu şekilde özetlenebilir.')
add_para('Birinci dönemde Selenium tabanlı web kazıma yöntemiyle oluşturulan TikTok veri seti, ikinci dönemde analitik dönüşüm süreçlerinden geçirilerek optimizasyon modeline hazır hale getirilmiştir. Etkileşim vekili hesaplamaları, takipçi sayısının tek başına performansı açıklamadaki yetersizliğini bir kez daha doğrulamıştır; yüksek takipçili bazı hesapların, düşük takipçili ama etkileşim odaklı hesapların gerisinde MCDM skoru aldığı gözlemlenmiştir.')
add_para('Gurobi ILP modeli, farklı bütçe senaryolarında (50K, 150K, 300K, 600K TL) tutarlı ve matematiksel olarak optimal portföyler üretmektedir. Model, geleneksel sezgisel yaklaşımların aksine, portföy bütçesini katmanlar arasında dengeleyerek hem maliyet etkinliğini hem de kategori çeşitliliğini optimize etmektedir.')
add_para('Çalışmanın sınırlılıkları değerlendirildiğinde, maliyet modelinin gerçek piyasa fiyatlarının sektör benchmarklarına dayanan bir tahmini olduğu ve bireysel pazarlık koşullarını yansıtmadığı görülmektedir. Veri seti yalnızca belirli bir zaman dilimini kapsamakta olup influencer performansı zaman içinde dalgalanabilmektedir. Web kazıma yöntemi yalnızca herkese açık profillerle sınırlıdır.')
add_para('Sonraki aşamada (Hafta 7-14) planlananlara bakıldığında, farklı bütçe seviyeleri ve kriter ağırlık kombinasyonları için kapsamlı duyarlılık analizi (Sensitivity Analysis) gerçekleştirilecektir. Modelin farklı marka profilleri için örnek vaka çalışmaları hazırlanacaktır. Streamlit arayüzüne senaryo karşılaştırma ve sonuç dışa aktarma özellikleri eklenecektir. Tüm çalışma, ENM458 tez yazım kılavuzuna uygun biçimde formatlanarak nihai forma taşınacaktır.')

add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# KAYNAKÇA
# ═══════════════════════════════════════════════════════════════════════════

add_heading('KAYNAKÇA', level=1)
refs = [
    '[1]\tJohnson, M., & Alvarez, L. (2021). The Rise of Digital Influencers: Shaping the Future of Marketing. Journal of Digital Communication, 14(3), 210–225.',
    '[2]\tKumar, R., & Lee, D. (2020). Influencer Marketing with Social Platforms. Social Media Research Journal, 8(4), 112–130.',
    '[3]\tZhang, P., & Han, Y. (2023). Social Media Influencer Marketing: Foundations, Trends and Research Directions. International Journal of Marketing Science, 19(1), 55–74.',
    '[4]\tAndersson, S., & Bergström, E. (2020). Instagram and Influencer Marketing: An Empirical Study of the Parameters Behind Success. Procedia Economics and Business, 7(2), 89–101.',
    '[5]\tDe Veirman, M., Cauberghe, V., & Hudders, L. (2017). Marketing through Instagram influencers: The impact of number of followers and product divergence. International Journal of Advertising, 36(5), 798–828.',
    '[6]\tLou, C., & Yuan, S. (2019). Influencer marketing: How message value and credibility affect consumer trust. Journal of Interactive Advertising, 19(1), 58–73.',
    '[7]\tSaito, M., & Kobayashi, K. (2019). A SI Model for Social Media Influencer Maximization. IEEE Access, 7, 150876–150889.',
    '[8]\tTiukhova, L., Korovin, D., & Melnikov, P. (2022). Influencer Prediction with Dynamic Graph Neural Networks. Neural Networks, 154, 145–159.',
    '[9]\tLi, C., & Wang, H. (2022). Computational Studies in Influencer Marketing. Expert Systems with Applications, 193, 116–127.',
    '[10]\tKim, S., & Cho, Y. (2023). Optimal Influencer Marketing Campaign under Budget Constraints. Journal of Business Analytics, 12(3), 180–195.',
    '[11]\tLópez-Dawn, A., & Giovanidis, A. (2021). Budgeted Portfolio Optimization Model for Social Media Influencer Selection. Journal of Applied Optimization, 18(4), 300–315.',
    '[12]\tAraujo, T., Neijens, P., & Vliegenthart, R. (2019). Discovering Effective Influencers. Computers in Human Behavior, 98, 10–20.',
    '[13]\tPhua, J., Jin, S. V., & Kim, J. (2020). The roles of celebrity endorsers\' credibility and attractiveness in influencer marketing. Computers in Human Behavior, 102, 310–321.',
    '[14]\tInfluencer Marketing Hub. (2023). Influencer Marketing Benchmark Report 2023. Global Industry Review Series.',
    '[15]\tDwivedi, Y. K., et al. (2021). Setting the future of digital and social media marketing research. International Journal of Information Management, 59, 102168.',
]
for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(ref)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# EKLER
# ═══════════════════════════════════════════════════════════════════════════

add_heading('EKLER', level=1)
add_heading('Tablo E.1. İş-Zaman Çizelgesi', level=2)

t5 = doc.add_table(rows=8, cols=4)
t5.style = 'Table Grid'
t5.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['No', 'İş Paketinin Adı', 'Zaman Aralığı', 'Takım Lideri']):
    set_cell_style(t5.rows[0].cells[i], h, bold=True, size=10, bg='D9D9D9')
jobs = [
    ('1', 'Veri Toplama Stratejisinin Belirlenmesi ve Ön Analiz',          'Ekim – Kasım 2025',         'Esra Zeyrek'),
    ('2', 'Kapsamlı Verinin Toplanması ve Veri Temizleme',                 'Kasım – Aralık 2025',       'A. Raif Yıldırım'),
    ('3', 'Verilerin İşlenmesi ve Keşifsel İstatistiksel Analiz',          'Aralık 2025 – Ocak 2026',   'Oğuzhan Dikmen'),
    ('4', 'Engagement Proxy, MCDM Skoru ve Maliyet Modeli Geliştirme',    'Ocak – Şubat 2026',         'A. Raif Yıldırım'),
    ('5', 'ILP Modelinin Kurulması, Gurobi Entegrasyonu ve Streamlit',     'Şubat – Mart 2026',         'Esra Zeyrek'),
    ('6', 'Senaryo Analizleri, Duyarlılık Testi ve Model Doğrulama',       'Nisan – Mayıs 2026',        'Oğuzhan Dikmen'),
    ('7', 'Sonuçların Değerlendirilmesi, Raporlama ve Tez Yazımı',         'Mayıs 2026',                'Esra Zeyrek'),
]
for r, row in enumerate(jobs):
    for c, val in enumerate(row):
        set_cell_style(t5.rows[r+1].cells[c], val, size=10)

doc.add_paragraph()
add_heading('Tablo E.2. Risk Yönetim Tablosu (B Planı)', level=2)

t6 = doc.add_table(rows=8, cols=3)
t6.style = 'Table Grid'
t6.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['No', 'En Büyük Riskler', 'Risk Yönetimi (B Planı)']):
    set_cell_style(t6.rows[0].cells[i], h, bold=True, size=10, bg='D9D9D9')
risks = [
    ('1', 'TikTok\'un veri erişimini kısıtlaması veya doğrulama mekanizmalarını artırması',
          'Esnek scraper yapısı kullanılmaktadır. Kısıtlama artarsa veri çekme hızı azaltılacak, mevcut veri seti ek kaynaklarla zenginleştirilecektir.'),
    ('2', 'Toplanan verinin analiz için yetersiz kalması',
          'Başlangıç örneklemi genişletilebilir biçimde tasarlanmıştır. Yetersizlik durumunda kullanıcı sayısı artırılacak veya normalize metrikler kullanılacaktır.'),
    ('3', 'API eksikliği nedeniyle veri tutarlılığının sağlanamaması',
          'Veriler yalnızca herkese açık kaynaklardan toplanmış, veri temizleme ve çapraz doğrulama ile tutarlılık güvence altına alınmıştır.'),
    ('4', 'Web kazıma sürecinde teknik hatalar veya kesintiler yaşanması',
          'Ara kayıt (checkpoint) mekanizmalarıyla veri kaybı önlenmiştir. Hata durumunda süreç kaldığı yerden devam eder.'),
    ('5', 'Optimizasyon modelinde uygulanabilir çözüm bulunamaması (infeasible)',
          'Bütçe artırılacak, kategori kısıtları esnetilecek veya kara liste daraltılacaktır. Arayüz bu durumda kullanıcıya öneride bulunmaktadır.'),
    ('6', 'Gurobi lisans sorunları yaşanması',
          'PuLP + CBC (açık kaynak) çözücüsüne geçiş B planı olarak hazır tutulmaktadır. Model bu geçişe uyumlu biçimde kodlanmıştır.'),
    ('7', 'Zaman planına uyulamaması',
          'İş paketleri paralel planlanmıştır. Gecikme durumunda kapsam daraltılarak temel model tutulacak, gelişmiş özellikler sonraki versiyona aktarılacaktır.'),
]
for r, row in enumerate(risks):
    for c, val in enumerate(row):
        set_cell_style(t6.rows[r+1].cells[c], val, size=10)

# ─── KAYDET ─────────────────────────────────────────────────────────────────
output_path = 'ENM458_Ara_Rapor_Final.docx'
doc.save(output_path)
print(f'✅ Dosya kaydedildi: {output_path}')
