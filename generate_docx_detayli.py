import docx
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = docx.Document()

# Adjust margins according to Tez Kılavuzu: Top 3, Bottom 2.5, Left 3.5, Right 2.5 (Cm)
section = doc.sections[0]
section.top_margin = Cm(3)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(3.5)
section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
paragraph_format = style.paragraph_format
paragraph_format.line_spacing = 1.5
paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = docx.shared.RGBColor(0, 0, 0)
        run.font.bold = True
    return h

def add_paragraph(text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1) 
    if bold_prefix:
        p.add_run(bold_prefix).bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

# Cover / Title Area
title = doc.add_paragraph()
title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("ENM458 ENDÜSTRİ MÜHENDİSLİĞİ BİTİRME PROJESİ II\nDETAYLI ARA SINAV RAPORU\n")
run.bold = True

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.LEFT
info.add_run("Tezin Adı: ").bold = True
info.add_run("TİKTOK’TA ETKİLİ MARKA İŞ BİRLİĞİ İÇİN İÇERİK ÜRETİCİSİ SEÇİMİNİN OPTİMİZASYONU\n")
info.add_run("Öğrenciler: ").bold = True
info.add_run("Esra ZEYREK, Abdullah Raif YILDIRIM, Oğuzhan DİKMEN\n")
info.add_run("Danışman: ").bold = True
info.add_run("Asst. Prof. Dr. Zeliha ERGÜL AYDIN\n")

# Section 1
add_heading("1. Proje Planının Gözden Geçirilmesi ve Kapsamı", level=1)
add_paragraph("Modern dijital pazarlama ekosisteminde kısa süreli formattaki içeriklerin tüketimi doğrusal bir ivmeyle artmakta, bu bağlamda TikTok platformu marka-tüketici etkileşiminin merkezinde yer almaktadır. Literatürde markaların içerik üreticilerini (influencer) seçerken sıklıkla takipçi sayısı gibi tekil metriklere veya tamamen insan sezgisine dayanan geleneksel 'deneme-yanılma' stratejilerine başvurdukları görülmektedir. Bu geleneksel yaklaşım, çoğu zaman yanlış hedef kitle eşleştirilmelerine ve sınırlı reklam bütçelerinin israf edilmesine yol açmaktadır.")
add_paragraph("Projemizin ilk aşaması olan ENM457 Endüstri Mühendisliği Bitirme Projesi I dersinde, optimizasyon modeline girdi sağlayacak temel veritabanını oluşturabilmek için birincil veri toplama süreçlerine odaklanılmıştır. Türkiye için TikTok API kısıtlamaları bulunduğundan, veri altyapısı özgün bir Selenium tabanlı web kazıma (web scraping) algoritması geliştirilerek inşa edilmiştir. Bu mekanizma sayesinde profil bilgileri, video bağlamları ve etkileşim metrikleri toplanarak statik dosyalar (Excel) halinde kaydedilmiştir.")
add_paragraph("Çalışmamızın ikinci ve asıl sistem geliştirme bölümü olan ENM458 dersi kapsamında ise, ilk dönemde toplanan bu ham metaverilerin (metadata) işlenmesi, matematiksel bir Karar Destek Sistemine dönüştürülmesi planlanmıştır. 1 ile 6. haftalar arasındaki süreç planlanan takvime uygun biçimde işletilmiş olup; verilerin ön işlemesi tamamlanmış, Gurobi çözücüsü kullanılarak 0-1 Tam Sayılı Doğrusal Programlama (ILP) modeli geliştirilmiş ve sistemin son müşteri olan pazarlama birimleri tarafından kullanılabilmesi amacıyla Python tabanlı bir Streamlit web arayüzü kurulmuştur. Bu adımlar, projenin tamamen analitik bir ürün (SaaS) haline gelme yolundaki hedeflerine B planına gerek kalmadan tam uyum sağladığını göstermektedir.")

# Section 2
add_heading("2. Sistem ve Problem Verilerinin Analitik Dönüşümü", level=1)
add_paragraph("İlk dönemde toplanan takipçi sayısı, toplam beğeni, video sayısı ve sektörel kategoriler gibi veriler doğrudan optimizasyon aşamasında kullanılamayacağı için çeşitli analitik işleme ve dönüşüm adımlarından geçirilmiş; her içerik üreticisi için karar algoritmalarında kullanılmak üzere sayısal değerlendirme parametreleri oluşturulmuştur.")

add_heading("2.1. Etkileşim Vekili (Engagement Proxy) Oluşturulması", level=2)
add_paragraph("Influencer'ların profillerini incelediğimizde sadece takipçi büyüklüğünün reklam başarısını garantilemediği belirlenmiştir. Geliştirdiğimiz algoritma ile her içerik üreticisinin video başına ortalama beğeni sayıları ve takipçi başına beğeni oranları hesaplanmıştır.")
add_paragraph("Algoritmamızda etkileşim performansını tek bir standardize değere taşımak için veri setindeki değerler %60 takipçi başına beğeni oranı, %40 ise ortalama video beğeni logaritması ağırlığıyla normalize (Min-Max) edilmiş ve Engagement Proxy (Etkileşim Vekili) skoru elde edilmiştir. Bu ağırlıklandırma, organik sadık kitlesi olan küçük içerik üreticilerine ve genel büyük kitleli üreticilere sistemin adil davranmasını sağlamıştır.")

add_heading("2.2. Maliyet Tahmin Modeli (Cost Estimation Model)", level=2)
add_paragraph("Platformda içerik üreticilerinin markalardan talep ettiği net fiyat bilgisi açık kaynaklardan direkt okunamamaktadır. Bunu çözmek üzere sektör araştırmaları ve uzman görüşlerine dayanan Adım Fonksiyonlu (Tier-based) Maliyet Tahmin Modeli kodlanmıştır:")
add_paragraph("Nano (0-10,000): Taban ücret= 1,500 TL (Her 1000 erişim = +150 TL)")
add_paragraph("Micro (10,000-100,000): Taban ücret= 8,000 TL (Her 1000 erişim = +80 TL)")
add_paragraph("Mid (100k-500k): Taban ücret= 35,000 TL (+35 TL)")
add_paragraph("Macro (500k-1M): Taban ücret= 90,000 TL (+20 TL)")
add_paragraph("Mega (1M+): Taban ücret= 200,000 TL (+8 TL)")
add_paragraph("Ayrıca, influencer'ların içerik tiplerine göre sektörel enflasyon etkenleri de maliyete çarpım olarak eklenmiştir (Güzellik ve Bakım = 1.15x, Moda = 1.10x, Yemek = 0.95x vb.). Bu ince ayarlar modelin maliyet verilerini son derece gerçekçi boyutlara indirgemiştir.")

add_heading("2.3. Çok Kriterli Karar Verme (MCDM) Skorlaması", level=2)
add_paragraph("Tüm performans dinamiklerini tek bir hedef fonksiyon ağırlığına toplayabilmek için 5 kriterli yapı kurulmuştur:")
add_paragraph("1. Etkileşim Skoru (Engagement Proxy): %35 ağırlık")
add_paragraph("2. Takipçi Sayısı Logaritması (Follower Base): %20 ağırlık")
add_paragraph("3. Takipçi Başına Beğeni Oranı: %20 ağırlık")
add_paragraph("4. Maliyet Etkinliği Logaritması (Cost Efficiency): %15 ağırlık")
add_paragraph("5. Video Hacmi (Aktif İçerik Durumu): %10 ağırlık")
add_paragraph("Böylece veri setindeki yüzlerce içerik üreticisinin her biri için Cazibe Skoru (MCDM Score) hesaplanmış ve veriler optimizasyon modeline girmeye hazır hale gelmiştir.")

# Section 3
add_heading("3. Probleme İlişkin Çözüm Modeli: Gurobi ILP Entegrasyonu", level=1)
add_paragraph("Seçilen içerik üreticisi sayılarının, kategorilerinin ve toplam harcamaların bir matematiksel düzen (optimizasyon) içerisinde bulunması Problemin doğası gereği 0-1 (Seçildi ya da Seçilmedi) mantığında çalışmaktadır. Problemin çözümü için Tam Sayılı Doğrusal Programlama (Integer Linear Programming) kurulmuştur. Sistem kapasitesi çok yüksek olduğundan akademik lisans kullanılarak Python Gurobi (gurobipy) çözücü motoru tercih edilmiştir.")

add_heading("3.1. Amaç Fonksiyonu ve Karar Değişkeni", level=2)
add_paragraph("İlgili i. içerik üreticisi kampanya için seçilirse Karar Değişkeni 1, aksi halde 0'dır. Amaç fonksiyonumuz, verilen bütçe sınırları dahilinde markanın pazarlama için edineceği toplam kazancı (MCDM skorlarının kümülatif miktarını) maksimize etmektir.")

add_heading("3.2. İş Dünyası (Gerçekçi) Kısıtları", level=2)
add_paragraph("Model sadece teorik değil, bir şirketin masasında karşılaşacağı tüm operasyonel kısıt şartlarını barındırır:")
add_paragraph("1. Toplam Bütçe Kısıtı: Seçilen kombinasyon hedef bütçeyi (Örn: 150.000 TL) kesinlikle aşamaz.")
add_paragraph("2. Çeşitlilik / Risk Yönetimi Kısıtı: Tüm para tek kişiye yatırılamaz, en az 3 kişiden oluşan portföyler gereklidir. Ayrıca enflasyon ve riski gidermek için hesaba maksimum kota sınırı verilmiştir.")
add_paragraph("3. Kategori Dağılımı (Segmentasyon): Kullanıcı arayüzden örneğin hem 'Fitness' hem 'Comedy' kategorisinden en az bir influencer olmasını kısıt (limit=1) olarak atayabilir.")
add_paragraph("4. Kara Liste (Blacklist) Toleransı: PR ve itibar sorunları nedeniyle markanın asla çalışmak istemediği kullanıcı profilleri arayüzden seçilir ve denklem anında O (Sıfır) kısıtı olarak kilitlenir. Algoritma bu şahısları değerlendirme matrisi dışına iter.")

add_heading("3.3. Streamlit ile Karar Destek Arayüzü", level=2)
add_paragraph("Yalnızca konsol çıktılarından ibaret olan akademik projeler ticari hayatta hızla geride kalmaktadır. Bu açığı kapatmak için Streamlit framework kullanılarak dinamik bir kullanıcı deneyimi web uygulaması oluşturulmuştur. Bu sistem sayesinde teknik bilgisi olmayan bir pazarlama direktörü arayüz çubukları (slider) vasıtasıyla 'Maksimum Bütçe', 'Zorunlu Kategoriler' ve 'Kara Liste Profillerini' seçerek arkadaki devasa Gurobi algoritmasını tetiklemekte; sonuçları, influencer fotoğrafları ve seçilen strateji matrisiyle ekranda saniyeler içinde okuyabilmektedir.")

# Section 4
add_heading("4. Projenin Girişimcilik ve Yenilikçilik Açısından Katkısı", level=1)
add_paragraph("Geleneksel reklam piyasasında ajanslar hangi influencer ile çalışmalıyız sorusuna sadece tecrübeler, ajans içi ağ bağlantıları (networking) ve makrolarla cevap vermektedir. Bu sistem bir insanın analiz kapasitesinin alamayacağı yüzlerce matematiksel senaryo iterasyonunu otomatize etmektedir.")
add_paragraph("Veri Demokrasisi ve Ürünleşme Yatkınlığı (SaaS Potansiyeli): Geliştirilen Streamlit aracı, açıkça ticari pazarlanabilir (commercialize) bir üründür. Pazarlama ajanslarına aylık abonelik modeliyle (B2B) sunulabilecek asgari geçerli (MVP) bir ürün altyapısı mevcuttur.")
add_paragraph("Gerçek Zamanlı Simülasyon: Rakiplerinin 'kara liste' opsiyonunu göz önünde bulundurarak dinamik olarak optimizasyon modeli anlık limitleri hesaplamakta, bütçe-ROI duyarlılıklarını pazarlama departmanlarına saniyeler içinde görsel metrik kartları halinde sağlamaktadır. Bu özellikleriyle projemiz salt bir araştırma olmaktan çıkıp yazılımsal katma değer üreten bir mühendislik start-up arketipi sergilemektedir.")

# Section 5
add_heading("5. Projenin BM Sürdürülebilir Kalkınma Amaçları Çerçevesindeki Etkisi", level=1)
add_paragraph("Endüstri mühendisliğinin doğası salt kâr maksimizasyonu değil, adil kalkınma ilkelerini gözetmektir. Bu projenin yapı taşları iki temel Birleşmiş Milletler Küresel Sürdürülebilirlik Amacı (SKA) çerçevesinde değer üretmektedir:")

add_heading("5.1. SKA 8 - İnsana Yakışır İş ve Ekonomik Büyüme", level=2)
add_paragraph("Sosyal medyada kazancın dağılımı piramit sistemindedir: Mega içerik üreticileri reklam pastasının tamamını almakta iken, Nano ve Micro düzeyde organik ve verimli performans sergileyen KOBİ odaklı üreticiler markalarca fark edilmemektedir. Optimizasyon modelimiz 'maliyet etkinliğini' ana formüllerinden biri yaptığı için, algoritma dev bütçe yutan 1 Mega hesap seçmek yerine, rekabet gücü olan yetenekli 15 Micro ve Nano üreticiyi hedef kitleye adaletli şekilde paylaştırabilmektedir. Bu sayede tabandaki içerik üreticilerinin gelir istihdamına katkı verilmiş olmaktadır.")

add_heading("5.2. SKA 12 - Sorumlu Üretim ve Tüketim", level=2)
add_paragraph("Ticari dünyada bilinçsiz harcanan medya reklam bütçeleri, şirketin sermayesinin israf edilmesidir. Kurulan analitik optimizasyon sistemi sadece parasal karşılığı değil, pazar etkisini (MCDM) maksimize eder. Marka verimsiz reklam alanlarında bütçeyi tüketmek yerine veriye dayalı hamleyi gerçekleştirir. Kuramsal açıdan kurumsal finans kaynaklarının savurganlıktan arındırılmış sorumlu bir tüketim mekanizmasına dönüşümünü sağlar.")

# Section 6
add_heading("6. Proje Beklentileri ve Takip Eden Süreç (Gelecek Planı)", level=1)
add_paragraph("Mevcut durumda ana belkemiği ve veri iskeleti kurulan Karar Destek aracı beklentileri aşan bir ilerleme performansı yakalamıştır. Takip eden proje altı haftasında (Hafta 7 - 14): Gurobi ILP senaryolarının kıyaslamalı duyarlılık analizlerinin yapılması, ağırlık faktörlerinin artıp azalması durumuyla ortaya çıkacak sonuç değişikliklerinin rapor edilmesi, arayüz kodlarının stabilitesinin test edilerek nihai tez yazım aşamasına geçilmesi ve tüm çıktının formata göre kitaplaştırılması hedeflenmektedir.")

doc.save('ENM458_Ara_Rapor_Detayli.docx')
print("ENM458_Ara_Rapor_Detayli.docx başarıyla oluşturuldu!")
