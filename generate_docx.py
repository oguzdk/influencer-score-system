import docx
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = docx.Document()

# Adjust margins according to Tez Kılavuzu: Top 3, Bottom 2.5, Left 3.5, Right 2.5 (Cm)
section = doc.sections[0]
section.top_margin = Cm(3)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(3.5)
section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
paragraph_format = style.paragraph_format
paragraph_format.line_spacing = 1.5
paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = docx.shared.RGBColor(0, 0, 0)
        run.font.bold = True
    return h

def add_paragraph(text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1) # Pragraflar 1cm içeriden başlar
    if bold_prefix:
        p.add_run(bold_prefix).bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

# Cover / Title Area
title = doc.add_paragraph()
title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("ENM458 ENDÜSTRİ MÜHENDİSLİĞİ BİTİRME PROJESİ II\nARA SINAV RAPORU\n")
run.bold = True

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.LEFT
info.add_run("Tezin Adı: ").bold = True
info.add_run("TİKTOK’TA ETKİLİ MARKA İŞ BİRLİĞİ İÇİN İÇERİK ÜRETİCİSİ SEÇİMİNİN OPTİMİZASYONU\n")
info.add_run("Öğrenciler: ").bold = True
info.add_run("Esra ZEYREK, Abdullah Raif YILDIRIM, Oğuzhan DİKMEN\n")
info.add_run("Danışman: ").bold = True
info.add_run("Asst. Prof. Dr. Zeliha ERGÜL AYDIN\n")

# Intro context
add_paragraph("Bu rapor, ENM458 Endüstri Mühendisliği Bitirme Projesi II kılavuzundaki 1-6. hafta gereksinimleri (Ara Sınav Raporu) doğrultusunda hazırlanmıştır.")

# Section 1
add_heading("1. Proje Planının Gözden Geçirilmesi ve Revizyonu", level=1)
add_paragraph("İlk dönemde (ENM457) TikTok platformundan içerik üreticilerine ait profil ve video metrikleri Selenium tabanlı web kazıma (web scraping) yöntemleriyle elde edilmiş ve ön işlemeleri tamamlanmıştır. İçerisinde bulunduğumuz ikinci dönemde (ENM458), elde edilen analitik veriler ışığında optimizasyon modelinin kurulması ve bu modelin kullanıcı dostu bir arayüzle buluşturulması hedeflenmiştir.")
add_paragraph("İlk altı haftalık süreçte proje planında hedeflenen terminlere uygun olarak ilerlenmiş; Gurobi Solver kullanılarak Tam Sayılı Doğrusal Programlama (ILP) modeli geliştirilmiş ve Streamlit kütüphanesi yardımıyla dinamik bir karar destek arayüzü kurulmuştur. Bu doğrultuda iş-zaman çizelgesine uygunluk sürmektedir ve B planına geçişi gerektirecek herhangi bir yöntem veya veri kısıtlaması yaşanmamıştır.")

# Section 2
add_heading("2. Sistem ve Problem Verilerinin Toplanması ile Analizinin Yapılması", level=1)
add_paragraph("İlk dönem toplanan TikTok metaverileri (takipçi sayısı, beğeni, video sayısı ve kategoriler), çalışmanın karar destek sistemi aşamasına hazırlanmak üzere ikinci dönemde çeşitli veri işleme ve dönüşüm adımlarından geçirilmiştir.")
p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(1)
p.add_run("• Etkileşim Vekili (Engagement Proxy): ").bold = True
p.add_run("İçerik üreticilerinin video başına ortalama beğeni sayıları ve takipçi sayıları üzerinden normalize edilmiş bir etkileşim göstergesi oluşturulmuştur.")

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(1)
p.add_run("• Tahmini Maliyetleme (Cost Estimation): ").bold = True
p.add_run("İçerik üreticileri takipçi bazlı (Nano, Micro, Mid, Macro, Mega) sınıflarına ayrılarak taban ücretler ve etkileşim başına oranlarla bütçe maliyet tahminleri yapılmıştır. Kategorilere bağlı çarpanlar (Beauty: 1.15, Food: 0.95 vb.) da tasarıma dahil edilerek maliyet asimetrileri gerçekçi koşullara uyarlanmıştır.")

p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(1)
p.add_run("• Çok Kriterli Karar Verme (MCDM) Skorlaması: ").bold = True
p.add_run("Etkileşim oranı, takipçi büyüklüğü, maliyet etkinliği ve video hacmi gibi öznitelikler ağırlıklandırılarak her içerik üreticisi için nihai bir 'MCDM (Karar) Skoru' üretilmiştir. Bu veriler model için hedef değişkeni oluşturmuştur.")

# Section 3
add_heading("3. Probleme İlişkin Çözüm Önerileri ve Seçilen Yöntem/Model", level=1)
add_paragraph("Karar verici olan markanın farklı bütçe kısıtları altında en iyi pazar etkisini (maksimum toplam MCDM skorunu) yakalayabilmesini sağlamak amacıyla deterministik yapısıyla esnek ve net sonuçlar üreten Tam Sayılı Doğrusal Programlama (Integer Linear Programming - ILP) modeli seçilmiştir.")
p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(1)
p.add_run("Amaç Fonksiyonu ve Kısıtlar: ").bold = True
p.add_run("Uygulanan model, seçilen içerik üreticilerinin kümülatif MCDM skorunu maksimize etmeyi amaçlar. Bunu yaparken markanın uygulayabileceği gerçekçi koşulları temsil eden kısıtlar sisteme entegre edilmiştir. Seçilenlerin toplam maliyetinin hedeflenen bütçe (B) sınırını aşmasına izin verilmez. Modele ayrıca minimum ve maksimum kişi seçimi kısıtları tanımlanarak marka portföyünde kampanya çeşitliliği sağlanmıştır. Kara Liste (Blacklist) mekanizması algoritmaya 0-1 kısıtı olarak entegre edilmiş, kullanıcının çalışmak istemediği hesaplar kesin olarak model dışı bırakılmıştır.")
p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(1)
p.add_run("Sistem Çözücü Motoru: ").bold = True
p.add_run("Büyük ölçekli kombinatoryal problemleri saniyeler (Time Limit toleransı) içerisinde çözebilme kapasitesine sahip olan ve lisanslı altyapısıyla güvenirlik sunan Gurobi Solver (gurobipy kütüphanesi) seçilmiştir. Oluşturulan Gurobi ILP modeli, Python Streamlit üzerinden bir gösterge paneli ve arayüz aracılığıyla kullanıcının parametrik girdilerine göre dinamik çözümler üretmektedir.")

# Section 4
add_heading("4. Projenin Girişimcilik ve Yenilikçilik Açısından Katkısı", level=1)
add_paragraph("Projenin yalnızca teorik bir akademik çalışma olarak kalmaması amacıyla kurulan Streamlit interaktif arayüzü, markalar ve dijital pazarlama ajansları için doğrudan pazarlanabilir bir ürün (SaaS) niteliği taşımaktadır. Günümüzde genellikle pazar sezgilerine ve sınırlı metrik analizlerine (örneğin sadece takipçi sayısına) dayanarak ilerleyen influencer pazarlaması, bu yazılım sayesinde tamamen veri güdümlü, analitik hesaplamalara dayalı bir karar destek sistemine dönüştürülmüştür. Markalara özgü kara liste uygulaması sağlaması, anlık bütçe sınırlarına ve istenen etiket kategorisi zorunluluklarına göre farklı senaryolar üretmesi projenin ticarileşebilir (inovatif) yönünü oluşturmaktadır.")

# Section 5
add_heading("5. Projenin BM Sürdürülebilir Kalkınma Amaçları Çerçevesindeki Etkileri", level=1)
add_paragraph("Geliştirilen karar destek modelinin teknolojik ve endüstriyel etkileri Birleşmiş Milletler Sürdürülebilir Kalkınma Amaçları (SKA) ile yakından ilişkilidir:")
p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(1)
p.add_run("SKA 8 (İnsana Yakışır İş ve Ekonomik Büyüme): ").bold = True
p.add_run("Geliştirilen matematiksel model, dijital reklam ekosisteminde aslan payını alan 'Mega' influencer'lar dışındaki 'Nano' ve 'Micro' seviyedeki yerel içerik üreticilerini de performans ve veri odağı ile sisteme çekebilmektedir. KOBİ'ler ve küçük reklamverenlerin de ulaşabileceği bütçe senaryoları ile istihdam çeşitliliği adil şekilde potansiyel işbirliklerine dahil edilmektedir.")
p = doc.add_paragraph()
p.paragraph_format.first_line_indent = Cm(1)
p.add_run("SKA 12 (Sorumlu Üretim ve Tüketim): ").bold = True
p.add_run("Model, dijital pazarlama kaynaklarını hedef odaklı (MCDM başarısı en yüksek olacak şekilde) değerlendirip israfları engellemektedir. Çıktıların bütçe performansına dönük olması sayesinde finansal anlamda gereksiz harcamaların, verimsiz tüketim alışkanlıklarının engellenmesi teşvik edilmekte, böylece ekonomik açıdan son derece sorumlu bir tüketim davranışı desteklenmektedir.")

doc.save('/Users/oguzhandikmen/Downloads/tezproje/ENM458_Ara_Rapor.docx')
