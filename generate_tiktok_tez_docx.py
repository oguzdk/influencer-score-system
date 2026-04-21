#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
ENM458 Ara Rapor - tiktok_influencer_tez.md → DOCX dönüştürücü
Tez Yazım Kılavuzuna uygun: Times New Roman 12pt, 1.5 satır, kenar boşlukları
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── SAYFA AYARLARI ────────────────────────────────────────────────────────
sec = doc.sections[0]
sec.page_width  = Cm(21)
sec.page_height = Cm(29.7)
sec.top_margin    = Cm(3.0)
sec.bottom_margin = Cm(2.5)
sec.left_margin   = Cm(3.5)
sec.right_margin  = Cm(2.5)

# ── NORMAL STİL ──────────────────────────────────────────────────────────
ns = doc.styles['Normal']
ns.font.name = 'Times New Roman'
ns.font.size = Pt(12)
ns.paragraph_format.space_before = Pt(0)
ns.paragraph_format.space_after  = Pt(6)
ns.paragraph_format.line_spacing = Pt(18)
ns.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.JUSTIFY

for i in range(1, 4):
    h = doc.styles[f'Heading {i}']
    h.font.name = 'Times New Roman'
    h.font.color.rgb = RGBColor(0,0,0)
    h.font.bold = True
    h.font.size = Pt(12)
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after  = Pt(3)
    h.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.LEFT

# ── YARDIMCILAR ──────────────────────────────────────────────────────────
def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name='Times New Roman'; r.font.color.rgb=RGBColor(0,0,0)
        r.font.bold=True; r.font.size=Pt(12)
    return h

def para(text='', indent=True, bold=False, italic=False,
         align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12):
    p = doc.add_paragraph()
    p.paragraph_format.alignment    = align
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    if indent: p.paragraph_format.first_line_indent = Cm(1)
    if text:
        r = p.add_run(text)
        r.font.name='Times New Roman'; r.font.size=Pt(size)
        r.bold=bold; r.italic=italic
    return p

def mixed(parts, indent=True, size=12):
    p = doc.add_paragraph()
    p.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    if indent: p.paragraph_format.first_line_indent = Cm(1)
    for text,bold,italic in parts:
        r = p.add_run(text)
        r.font.name='Times New Roman'; r.font.size=Pt(size)
        r.bold=bold; r.italic=italic
    return p

def cell_style(cell, text, bold=False, size=10, bg=None):
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.font.name='Times New Roman'; r.font.size=Pt(size); r.bold=bold
    if bg:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto')
        shd.set(qn('w:fill'), bg)
        tcPr.append(shd)

def tbl_caption(text):
    p = doc.add_paragraph()
    p.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run(text)
    r.font.name='Times New Roman'; r.font.size=Pt(10); r.bold=True

def space(): doc.add_paragraph()

def pb(): doc.add_page_break()

# ════════════════════════════════════════════════════════════════════
# KAPAK
# ════════════════════════════════════════════════════════════════════
for txt, bold in [
    ('TİKTOK\'TA ETKİLİ MARKA İŞ BİRLİĞİ İÇİN\nİÇERİK ÜRETİCİSİ SEÇİMİNİN OPTİMİZASYONU', True),
    ('Lisans Tezi', False),
]:
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(txt); r.font.name='Times New Roman'; r.font.size=Pt(12); r.bold=bold

space()
p = doc.add_paragraph()
p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Esra ZEYREK\nAbdullah Raif YILDIRIM\nOğuzhan DİKMEN')
r.font.name='Times New Roman'; r.font.size=Pt(12)

space(); space()
for txt, bold in [
    ('ENM458 ENDÜSTRİ MÜHENDİSLİĞİ BİTİRME PROJESİ II', True),
    ('ARA SINAV RAPORU', True),
]:
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(txt); r.font.name='Times New Roman'; r.font.size=Pt(12); r.bold=bold

space()
for txt in [
    'Danışman: Asst. Prof. Dr. Zeliha ERGÜL AYDIN',
    'Eskişehir Teknik Üniversitesi Mühendislik Fakültesi\nEndüstri Mühendisliği Bölümü\nNisan, 2026',
]:
    p = doc.add_paragraph()
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(txt); r.font.name='Times New Roman'; r.font.size=Pt(12)

pb()

# ════════════════════════════════════════════════════════════════════
# ÖZET
# ════════════════════════════════════════════════════════════════════
heading('ÖZET', 1)
para('TİKTOK\'TA ETKİLİ MARKA İŞ BİRLİĞİ İÇİN İÇERİK ÜRETİCİSİ SEÇİMİNİN OPTİMİZASYONU', indent=False, bold=True)
para('Esra ZEYREK | Abdullah Raif YILDIRIM | Oğuzhan DİKMEN', indent=False)
para('Endüstri Mühendisliği — Eskişehir Teknik Üniversitesi, Mühendislik Fakültesi, Nisan, 2026\nDanışman: Asst. Prof. Dr. Zeliha ERGÜL AYDIN', indent=False)
para('Dijital pazarlama faaliyetlerinin sosyal medya platformlarına kaymasıyla birlikte, içerik üreticileri markalar için önemli bir iletişim aracı hâline gelmiştir. Özellikle TikTok, kısa video formatı ve yüksek etkileşim oranları sayesinde marka iş birliklerinin yoğunlaştığı bir platform olarak öne çıkmaktadır. Ancak içerik üreticisi (influencer) seçimi çoğu zaman takipçi sayısı gibi sınırlı göstergelere dayandırılmakta, bu durum reklam bütçelerinin etkin kullanılmasını zorlaştırmaktadır.')
para('Bu çalışma, TikTok üzerinde gerçekleştirilecek marka iş birlikleri için veri temelli ve matematiksel bir içerik üreticisi seçim modeli geliştirmeyi amaçlamaktadır. ENM457 (Bitirme Projesi I) kapsamında Selenium tabanlı web kazıma yöntemiyle içerik üreticilerine ait profil ve video metrikleri sistematik biçimde toplanmış; bu veriler ön işleme ve keşifsel analizden geçirilmiştir. ENM458 (Bitirme Projesi II) kapsamında ise her içerik üreticisi için etkileşim vekili (engagement proxy) hesaplanmış, takipçi kademeleri ve kategori çarpanlarına dayanan bir maliyet tahmin modeli oluşturulmuş ve çok kriterli karar verme (MCDM) skoru üretilmiştir. Bu sayısal göstergeler, Gurobi akademik lisansıyla çözülen 0-1 Tam Sayılı Doğrusal Programlama (ILP) modeline girdi olarak aktarılmış; model bütçe kısıtı, minimum portföy çeşitliliği, kategori dengesi ve kara liste toleransı gibi gerçekçi iş koşullarını barındırmaktadır. Sistem Streamlit tabanlı interaktif bir karar destek arayüzü ile son kullanıcıya sunulmaktadır.')
para('Elde edilen bulgular, modelin farklı bütçe senaryolarında tutarlı ve anlamlı influencer portföyleri önerdiğini, takipçi büyüklüğü yerine etkileşim kalitesine dayalı çözümler ürettiğini ortaya koymaktadır. Yaklaşım; dijital reklam kaynaklarının daha verimli kullanılmasına katkı sağlamakta ve BM SKA 8: İnsana Yakışır İş ve Ekonomik Büyüme ile SKA 12: Sorumlu Üretim ve Tüketim hedefleriyle ilişkilendirilmektedir.')
mixed([('Anahtar Sözcükler: ', True, False), ('TikTok, içerik üreticisi seçimi, influencer optimizasyonu, 0-1 tam sayılı doğrusal programlama, MCDM, Gurobi, web kazıma, dijital pazarlama, karar destek sistemi, sorumlu tüketim', False, False)], indent=False)
pb()

# ════════════════════════════════════════════════════════════════════
# ABSTRACT
# ════════════════════════════════════════════════════════════════════
heading('ABSTRACT', 1)
para('OPTIMIZATION OF CONTENT CREATOR SELECTION FOR EFFECTIVE BRAND COLLABORATIONS ON TIKTOK', indent=False, bold=True)
para('Esra ZEYREK | Abdullah Raif YILDIRIM | Oğuzhan DİKMEN', indent=False)
para('Department of Industrial Engineering — Eskisehir Technical University, Engineering Faculty, April, 2026\nSupervisor: Asst. Prof. Dr. Zeliha ERGÜL AYDIN', indent=False)
para('With the rapid expansion of social media platforms, content creators have become a central element of brand communication strategies. TikTok has emerged as a dominant platform for brand collaborations due to its short-video format and high user engagement rates. However, influencer selection is frequently based on oversimplified indicators such as follower count, leading to inefficient use of advertising budgets and poor audience alignment.')
para('This study develops a data-driven, mathematically grounded content creator selection model for TikTok-based brand collaborations. In the first phase (ENM457), profile and video-level engagement data were systematically collected via Selenium-based web scraping and subjected to exploratory analysis. In the second phase (ENM458), an engagement proxy was computed for each creator, a tiered cost estimation model was constructed using follower brackets and category multipliers, and a Multi-Criteria Decision Making (MCDM) score was derived. These indicators serve as inputs to a 0-1 Integer Linear Programming (ILP) model solved with Gurobi, incorporating realistic business constraints including budget limits, minimum portfolio diversity, category distribution requirements, and blacklist filtering. The system is delivered through an interactive Streamlit-based decision support interface.')
para('Results demonstrate that the model consistently produces meaningful influencer portfolio recommendations across diverse budget scenarios, prioritizing engagement quality over follower volume. The approach directly aligns with SDG 8: Decent Work and Economic Growth and SDG 12: Responsible Consumption and Production.')
mixed([('Keywords: ', True, False), ('TikTok, content creator selection, influencer optimization, 0-1 integer linear programming, MCDM, Gurobi, web scraping, digital marketing, decision support system, responsible consumption', False, False)], indent=False)
pb()

# ════════════════════════════════════════════════════════════════════
# İÇİNDEKİLER
# ════════════════════════════════════════════════════════════════════
heading('İÇİNDEKİLER', 1)
items = [
    ('1. GİRİŞ', False),
    ('2. PROBLEMİN TARİFİ', False),
    ('3. LİTERATÜR TARAMASI', False),
    ('    3.1. Genel Değerlendirme ve Tezin Literatürdeki Konumu', False),
    ('4. PROJENİN GİRİŞİMCİLİK VE YENİLİKÇİLİK AÇISINDAN KATKISI', False),
    ('    4.1. Projenin BM SKA Kapsamında Topluma, Ekonomiye, Sürdürülebilirliğe ve Çevreye Etkileri', False),
    ('5. ÇALIŞMADA KULLANILAN TEMEL KAVRAMLAR VE YÖNTEMLER', False),
    ('    5.1. Veri Toplama Süreci ve Ön İşleme', False),
    ('    5.2. Analitik Dönüşüm Adımları', False),
    ('    5.3. Optimizasyon Modeli: 0-1 Tam Sayılı Doğrusal Programlama', False),
    ('    5.4. Streamlit Tabanlı Karar Destek Arayüzü', False),
    ('6. SONUÇ VE ÖNERİLER', False),
    ('KAYNAKÇA', False),
    ('7. EKLER', False),
]
for txt, bold in items:
    p = doc.add_paragraph()
    p.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.first_line_indent = Cm(0)
    r = p.add_run(txt); r.font.name='Times New Roman'; r.font.size=Pt(12); r.bold=bold

space()
heading('ŞEKİLLER DİZİNİ', 2)
for s in ['Şekil 1: TikTok İçerik Üreticisi Seçimi ve Optimizasyonu Genel İş Akış Şeması',
          'Şekil 2: Takipçi Sayısı vs. Beğeni Sayısı İlişkisi',
          'Şekil 3: Video Süresi vs. Beğeni Sayısı',
          'Şekil 4: Takipçi Sayısı vs. İçerik Üreticisi Skoru']:
    para(s, indent=False, size=10)

space()
heading('TABLOLAR DİZİNİ', 2)
for t in ['Tablo 5.1. Takipçi Bazlı Kademeli Maliyet Modeli',
          'Tablo 5.2. Kategori Maliyet Çarpanları',
          'Tablo 5.3. MCDM Kriter Ağırlık Tablosu',
          'Tablo 5.4. Gurobi Çözücü Parametre Tablosu',
          'Tablo E.1. İş-Zaman Çizelgesi',
          'Tablo E.2. Risk Yönetim Tablosu (B Planı)']:
    para(t, indent=False, size=10)

pb()

# ════════════════════════════════════════════════════════════════════
# 1. GİRİŞ
# ════════════════════════════════════════════════════════════════════
heading('1. GİRİŞ', 1)
para('Dijitalleşmenin hız kazanmasıyla birlikte pazarlama faaliyetleri önemli ölçüde dönüşmüş, markalar hedef kitlelerine ulaşmak için sosyal medya platformlarını daha yoğun biçimde kullanmaya başlamıştır. Özellikle kısa video temelli içeriklerin yaygınlaşması, kullanıcıların içerik tüketim alışkanlıklarını değiştirmiş ve markalar açısından yeni iletişim olanakları yaratmıştır. Bu bağlamda TikTok, algoritmik keşif yapısı ve yüksek etkileşim oranları sayesinde dijital pazarlama stratejilerinde öne çıkan platformlardan biri hâline gelmiştir.')
para('TikTok üzerinde yürütülen marka iş birliklerinde içerik üreticileri, yalnızca tanıtım yapan kişiler olarak değil, marka algısını şekillendiren önemli paydaşlar olarak değerlendirilmektedir. Ancak uygulamada içerik üreticisi seçimi çoğu zaman takipçi sayısı, toplam beğeni miktarı veya önceki iş birliklerine ilişkin genel izlenimlere dayalı olarak yapılmaktadır. Bu tür yaklaşımlar, içerik üreticisinin hedef kitleyle olan uyumunu ve gerçek etkileşim performansını yeterince yansıtmayabilmektedir.')
para('ENM457 (Bitirme Projesi I) kapsamında gerçekleştirilen ilk aşamada, TikTok platformundan içerik üreticilerine ait veri seti web kazıma yöntemleri ile elde edilmiş ve bu veriler ön incelemeye tabi tutulmuştur. Bu süreçte elde edilen veri altyapısı, çalışmanın ikinci aşaması için zemin oluşturmuştur.')
para('ENM458 (Bitirme Projesi II) kapsamındaki bu raporda ise, elde edilen veri seti ileri düzey analizler için dönüştürülmüş, içerik üreticisi seçimine yönelik çok kriterli değerlendirme modeli oluşturulmuş ve 0-1 Tam Sayılı Doğrusal Programlama (ILP) yöntemiyle optimal influencer portföyü belirlenmektedir. Geliştirilen sistem, Streamlit tabanlı bir karar destek arayüzüyle son kullanıcıya sunulmakta; teknik bilgisi olmayan pazarlama yöneticilerinin bile saniyeler içinde optimize edilmiş influencer önerisi almasını mümkün kılmaktadır.')
para('Bu rapor, ENM458 dersi 1 ile 6. haftalar arasındaki çalışmaları (Ara Sınav Raporu) kapsamakta olup modelin dayandığı yöntemi, veri hazırlık sürecini ve gerçekleştirilen uygulamayı bütünleşik biçimde aktarmaktadır.')

# ════════════════════════════════════════════════════════════════════
# 2. PROBLEMİN TARİFİ
# ════════════════════════════════════════════════════════════════════
heading('2. PROBLEMİN TARİFİ', 1)
para('Bu çalışma, dijital pazarlama alanında TikTok platformu üzerinde gerçekleştirilen marka–içerik üreticisi iş birliklerine odaklanmaktadır. TikTok, kısa video formatı ve yüksek etkileşim yapısı sayesinde markaların geniş kitlelere ulaşmasını mümkün kılmakta, bu durum içerik üreticisi seçiminin stratejik önemini artırmaktadır.')
para('Çalışmanın temel problemi, TikTok üzerinde gerçekleştirilecek marka iş birlikleri için uygun içerik üreticilerinin nasıl belirleneceğidir. Mevcut uygulamalarda bu süreç çoğunlukla takipçi sayısı veya öznel değerlendirmelere dayandırılmakta; video bazlı etkileşim yapıları ve içerik performansı yeterince dikkate alınmamaktadır. Bu durum, özellikle sınırlı bütçeye sahip kampanyalarda reklam kaynaklarının verimsiz kullanılmasına yol açmaktadır.')
para('TikTok\'un algoritmik yapısı, içerik üreticisi performansının yalnızca basit göstergelerle açıklanmasını zorlaştırmaktadır. Yüksek takipçi sayısı her zaman yüksek etki anlamına gelmemekte; etkileşim oranı, içerik türü ve kullanıcı davranışları daha belirleyici hâle gelmektedir. Türkiye\'de TikTok verilerine doğrudan erişim sağlayan bir API bulunmaması, veri temelli analizleri zorlaştırmakta ve alternatif veri toplama yöntemlerini gerekli kılmaktadır.')
para('Bu çalışma kapsamında problem iki aşamalı olarak ele alınmaktadır. İlk Aşamada (ENM457) içerik üreticilerine ait veriler web kazıma yöntemiyle toplanmış, veri seti temizlenmiş ve yalnızca Türkiye\'de faaliyet gösteren bireysel içerik üreticileri analize dahil edilmiştir. İkinci Aşamada (ENM458) ise veri seti analitik dönüşümden geçirilmiş; etkileşim vekili, maliyet tahmini ve MCDM skoru hesaplanmış, ardından belirli bir bütçe altında en yüksek etkiyi sağlayacak içerik üreticisi kombinasyonunu seçen bir 0-1 ILP optimizasyon modeli geliştirilmiş ve Streamlit arayüzüyle bütünleştirilmiştir.')
para('Problemin gerçekçi kısıtları şunlardır: ekonomik açıdan reklam bütçesi sınırlıdır ve her influencer\'ın tahmini maliyeti farklıdır; stratejik açıdan marka, belirli kategorilerde mutlaka temsilci istemektedir; itibar yönetimi açısından bazı influencer\'larla çalışılmak istenmemektedir (kara liste); risk yönetimi açısından ise tüm bütçenin tek bir hesaba ayrılması hem finansal hem de PR riski taşımaktadır.')

# ════════════════════════════════════════════════════════════════════
# 3. LİTERATÜR
# ════════════════════════════════════════════════════════════════════
heading('3. LİTERATÜR TARAMASI', 1)
para('Bu bölümde, çalışmanın konusunu oluşturan içerik üreticisi pazarlaması, içerik üreticisi performansının değerlendirilmesi, veri temelli ve hesaplamalı yaklaşımlar ile bütçe kısıtlı karar verme modelleri literatürdeki çalışmalar çerçevesinde ele alınmaktadır. İncelenen çalışmalar, içerik üreticisi seçiminin hangi kriterler doğrultusunda yapıldığını ortaya koymakta ve bu tez çalışmasının literatürde konumlandığı noktayı belirlemektedir.')
para('Dijital içerik üreticisi kavramı, sosyal medya platformlarının yaygınlaşmasıyla birlikte pazarlama literatüründe önemli bir yer edinmiştir. Johnson ve Alvarez, dijital içerik üreticilerinin markalar ile tüketiciler arasında güven temelli bir köprü kurduğunu ve geleneksel reklamcılıktan farklı bir etkileşim yapısı sunduğunu belirtmektedir [1]. Benzer şekilde Kumar ve Lee, içerik üreticisi pazarlamasının kullanıcı etkileşimi ve içerik özgünlüğü üzerine kurulu yapısını incelemiştir [2].')
para('Zhang ve Han, influencer pazarlamasının teorik temellerini, güncel eğilimlerini ve gelecekteki araştırma yönlerini kapsamlı biçimde ele almaktadır; sektör hacminin 2023 itibarıyla 21 milyar doları aştığı ve içerik üreticisi seçiminin giderek daha analitik bir yapıya kavuştuğu ifade edilmektedir [3]. Andersson ve Bergström ise ampirik bulgularla, influencer başarısının yalnızca takipçi sayısıyla ölçülemeyeceğini; içerik kalitesi ve etkileşim yapısının da belirleyici olduğunu ortaya koymuştur [4].')
para('Son yıllarda içerik üreticisi pazarlamasında hesaplamalı yöntemlerin kullanımı artış göstermiştir. Li ve Wang, büyük veri analitiği ve ağ tabanlı yöntemlerin içerik üreticisi etkinliğini anlamada önemli katkılar sunduğunu belirtmiştir [5]. Araujo, Neijens ve Vliegenthart ise etkili içerik üreticilerinin keşfedilmesinde ağ yapıları ve etkileşim örüntülerinin kritik rol oynadığını ortaya koymuştur [6].')
para('Saito ve Kobayashi tarafından geliştirilen SI modeli, içerik üreticisi maksimizasyon problemini matematiksel bir çerçevede ele almakta ve sosyal medya yayılımının analitik olarak modellenebileceğini göstermektedir [7]. Tiukhova, Korovin ve Melnikov tarafından önerilen dinamik grafik sinir ağları tabanlı yaklaşım ise içerik üreticisi tahminini zaman boyutunu da kapsayacak biçimde ele almış; statik değil dinamik değerlendirmenin önemini vurgulamıştır [8].')
para('Influencer Marketing Benchmark Report 2023, markaların influencer pazarlamasına ayırdığı bütçelerin arttığını, ancak bu bütçelerin etkin kullanımına yönelik analitik yaklaşımların hâlâ sınırlı olduğunu ortaya koymaktadır [9]. Kim ve Cho, bütçe kısıtları altında optimal influencer kampanyasının nasıl oluşturulabileceğini incelemiş ve veri temelli planlamanın kampanya başarısını anlamlı biçimde artırdığını göstermiştir [10]. López-Dawn ve Giovanidis ise içerik üreticisi seçimini bütçe kısıtlı bir portföy optimizasyon problemi olarak ele alarak sistematik bir çözüm önermiştir [11].')
para('De Veirman, Cauberghe ve Hudders, yüksek takipçi sayısının her zaman daha yüksek etki anlamına gelmediğini ve güven algısıyla doğrusal olmayan bir ilişki içinde olduğunu göstermiştir [12]. Lou ve Yuan ise influencer mesajlarının algılanan değeri ve güvenilirliğinin tüketici güvenini doğrudan etkilediğini belirtmiştir [13]. Phua, Jin ve Kim tarafından gerçekleştirilen çalışmada, parasosyal etkileşimin kullanıcı davranışları üzerindeki belirleyici rolü ortaya konmuştur [14].')
para('Dwivedi ve arkadaşları, dijital ve sosyal medya pazarlamasındaki geleceğe yönelik araştırma yönlerini incelerken bu tür karar problemlerinde hesaplamalı yöntemler, büyük veri analitiği ve karar destek modellerinin belirleyici rol oynayacağını vurgulamaktadır [15]. Mevcut tez çalışması, literatürde önerilen bu yönelimle tam uyum içinde biçimde, içerik üreticisi seçimini veri temelli ve optimizasyon odaklı bir çerçevede ele almaktadır.')

heading('3.1. Genel Değerlendirme ve Tezin Literatürdeki Konumu', 2)
para('Literatür incelendiğinde, içerik üreticisi pazarlamasının kavramsal temellerinin büyük ölçüde oluşturulduğu, ancak platforma özgü ve veri temelli çalışmaların hâlâ sınırlı olduğu görülmektedir. Özellikle TikTok gibi algoritmik keşif yapısına sahip platformlar için geliştirilen içerik üreticisi seçim yaklaşımları oldukça kısıtlıdır; mevcut çalışmalar büyük çoğunlukla Instagram veya YouTube gibi platformlarda yoğunlaşmakta, Türkiye ölçeğindeki TikTok dinamiklerini yansıtmamaktadır.')
para('Bu tez çalışması; TikTok özelinde gerçek platform verilerine dayalı veri toplama, çok kriterli skorlama ve 0-1 ILP optimizasyonunun bir arada kullanıldığı bütünleşik bir sistem sunarak literatürdeki önemli bir boşluğu doldurmaktadır. Çalışma, hem kuramsal hem de uygulamaya yönelik katkı sağlamaktadır: kuramsal katkı, TikTok influencer seçiminin matematiksel bir optimizasyon problemi olarak modellenmesiyle; uygulamaya yönelik katkı ise Streamlit arayüzü ile doğrudan kullanılabilir bir karar destek sistemi sunulmasıyla gerçekleştirilmektedir.')

pb()

# ════════════════════════════════════════════════════════════════════
# 4. GİRİŞİMCİLİK
# ════════════════════════════════════════════════════════════════════
heading('4. PROJENİN GİRİŞİMCİLİK VE YENİLİKÇİLİK AÇISINDAN KATKISI', 1)
para('Bu çalışma, dijital pazarlama alanında giderek artan influencer kullanımına rağmen influencer seçim sürecinin büyük ölçüde sezgisel ve deneyime dayalı biçimde yürütülmesine yönelik bir problemi ele almaktadır. Özellikle küçük ve orta ölçekli işletmeler açısından yanlış influencer tercihlerinin reklam bütçeleri üzerinde önemli kayıplara yol açabildiği görülmektedir. Bu bağlamda çalışma, influencer seçimini veri temelli ve analitik bir yaklaşımla ele alarak girişimcilik açısından uygulanabilir bir çözüm sunmayı hedeflemektedir.')
para('Girişimcilik perspektifinden değerlendirildiğinde, geliştirilen Streamlit arayüzü aylık abonelik modeliyle (B2B SaaS) dijital ajanslar ve markalara sunulabilecek, minimum uygulanabilir ürün (MVP) niteliğinde bir altyapı oluşturmaktadır. TikTok\'a özgü verilerin sistematik biçimde toplanması, temizlenmesi, analiz edilmesi ve matematiksel modele entegre edilmesiyle ortaya çıkan bu sistem; influencer seçim sürecinin ölçülebilir, karşılaştırılabilir ve optimize edilebilir hâle getirilmesine olanak tanımaktadır. Öte yandan arayüz, anlık bütçe kısıtlarına ve istenen kategori zorunluluklarına göre farklı senaryolar üretebilmekte; rakiplerinin kara liste opsiyonunu göz önüne alarak dinamik optimizasyon yapabilmektedir.')
para('Yenilikçilik açısından çalışmanın öne çıkan boyutları şunlardır: Türkiye\'de TikTok verisine dayalı influencer seçimi için geliştirilen ilk optimizasyon modeli olma özelliği; API erişimi gerektirmeyen, web kazıma tabanlı özgün veri toplama altyapısı; takipçi sayısı yerine etkileşim kalitesini merkeze alan MCDM skorlama sistemi; gerçek zamanlı senaryo simülasyonu; teknik bilgi gerektirmeyen sezgisel kullanıcı arayüzü. Proje aynı zamanda Endüstri Mühendisliği (optimizasyon modeli), Bilgisayar Mühendisliği (web kazıma ve Streamlit) ve İşletme (pazarlama stratejisi, maliyet modeli) disiplinlerini başarıyla entegre eden multidisipliner bir yapıya sahiptir.')
para('Geliştirilen yaklaşım, ilerleyen aşamalarda ticarileştirilebilir bir karar destek sistemine dönüştürülebilecek potansiyel taşımaktadır. Bu yönüyle proje, salt bir akademik çalışmanın ötesine geçerek yazılımsal katma değer üreten bir mühendislik girişim prototipi sergilemektedir.')

heading('4.1. Projenin BM Sürdürülebilir Kalkınma Amaçları Kapsamında Topluma, Ekonomiye, Sürdürülebilirliğe ve Çevreye Etkileri', 2)
para('Bu çalışma, Birleşmiş Milletler Sürdürülebilir Kalkınma Amaçları (SKA) kapsamında iki temel hedefle doğrudan ilişkilendirilmektedir.')
para('SKA 8 — İnsana Yakışır İş ve Ekonomik Büyüme: Dijital reklam dünyasındaki mevcut yapıda, harcamaların büyük bölümü Mega (1 milyondan fazla takipçili) içerik üreticilerinde yoğunlaşmaktadır. Bu durum, organik etkileşim oranları yüksek ancak henüz geniş kitlelere ulaşamamış Nano ve Micro düzeydeki yaratıcıların ekosistemin dışında kalmasına neden olmaktadır. Geliştirilen optimizasyon modeli, maliyet etkinliği kriterini sisteme dahil ederek bu dengesizliği düzeltmekte; KOBİ\'ler ve yerel markalar da uygun maliyetli ama etkili işbirliği yapma imkânı bulmaktadır.')
para('SKA 12 — Sorumlu Üretim ve Tüketim: Sezgisel kararlarla tahsis edilen reklam bütçeleri kaynak israfına ve başarısız kampanyaların tekrarlanmasına yol açmaktadır. Bu proje, her harcanan reklam lirasının matematiksel hesap görmeyi zorunlu kılan bir sistem kurarak sorumlu tüketimi teşvik etmekte ve dijital reklam sektöründe sürdürülebilir bir karar verme kültürünün yerleşmesine katkı sağlamaktadır.')
para('Toplumsal açıdan, daha hedefli içerik üretiminin teşvik edilmesi kullanıcıların ilgisiz reklamlara maruz kalmasını azaltarak dijital içerik tüketiminde daha dengeli bir yapı oluşturmaktadır. Ekonomik açıdan ise reklam bütçelerinin veri temelli kararlarla yönetilmesi, yatırım geri dönüşünü artırarak özellikle sınırlı kaynaklara sahip işletmeler için rekabet avantajı yaratmaktadır. Sürdürülebilirlik ve çevresel etkiler açısından, başarısız kampanyaların azaltılması gereksiz dijital içerik üretimini ve bunun getirdiği dolaylı enerji tüketimini de azaltmaktadır.')

pb()

# ════════════════════════════════════════════════════════════════════
# 5. YONTEM
# ════════════════════════════════════════════════════════════════════
heading('5. ÇALIŞMADA KULLANILAN TEMEL KAVRAMLAR VE YÖNTEMLER', 1)
para('Sosyal medya üzerinde yürütülen influencer pazarlama faaliyetlerinin etkinliği, içerik üreticilerine ait güvenilir, tutarlı ve çok boyutlu verilerin sistematik biçimde analiz edilebilmesine bağlıdır. Takipçi sayısı, etkileşim oranı, video başına beğeni ortalaması ve içerik üretim sıklığı gibi nicel göstergeler; içerik üreticilerinin marka açısından yaratacağı potansiyel etkiyi ölçmede belirleyici rol oynamaktadır.')
para('Ancak TikTok platformu, Türkiye özelinde resmi bir API sunmamaktadır. Bu durum, tüm veri toplama sürecini teknik ve operasyonel açıdan çalışmanın en kritik aşamalarından biri hâline getirmektedir. Bu çalışmada benimsenen metodolojik çerçeve; veri toplama ve ön işleme, analitik dönüşüm adımları, 0-1 ILP optimizasyon modeli ve Streamlit karar destek arayüzü olmak üzere dört ana aşamadan oluşmaktadır.')

heading('5.1. Veri Toplama Süreci ve Ön İşleme', 2)
para('İçerik üreticilerine ait veriler, Python programlama dili kullanılarak geliştirilen Selenium tabanlı dinamik web kazıma yöntemiyle elde edilmiştir. Veri toplama sürecine başlanmadan önce, analiz edilecek içerik üreticileri için bir başlangıç örneklemi oluşturulmuş; birden fazla kez yer alan kullanıcı adları temizlenmiş, kurumsal hesaplar ve içerik üreticisi pazarlaması kapsamında değerlendirilemeyecek profiller elenmiştir.')
para('Kazıma süreci kapsamında her içerik üreticisi için profil sayfasına erişilerek şu veriler toplanmıştır: kullanıcı adı, görünen ad, biyografi bilgisi, takipçi sayısı, takip edilen hesap sayısı, toplam beğeni, video sayısı. Video sayfalarından ise görüntülenme sayısı, beğeni, yorum ve paylaşım sayıları ile içerik kategorisi bilgileri elde edilmiştir. TikTok\'un doğrulama mekanizmaları devreye girdiğinde etik ve teknik sınırlar gözetilerek otomatik aşma yöntemleri kullanılmamış; manuel müdahale sonrası süreç kaldığı yerden sürdürülmüştür.')
para('Keşifsel analizler kapsamında, takipçi sayıları ile video başına beğeni sayıları arasındaki ilişki incelenmiş; elde edilen sonuçlar takipçi sayısının artmasının her durumda doğrusal etkileşim artışıyla sonuçlanmadığını ortaya koymaktadır. Video süresi ile beğeni sayısı arasındaki ilişki logaritmik ölçekte analiz edilmiş; farklı video sürelerine sahip içeriklerin benzer düzeylerde etkileşim elde edebildiği gözlemlenmiştir. Çok boyutlu performans skoru ile takipçi sayısı karşılaştırıldığında ise yüksek takipçi sayısının her zaman yüksek performans skoruna karşılık gelmediği görülmüştür.')

heading('5.2. Analitik Dönüşüm Adımları', 2)
para('Ham veri doğrudan optimizasyon modeline aktarılamaz. ENM458 kapsamında toplanan veriler üç aşamalı bir analitik dönüşüm sürecinden geçirilmiştir.')
para('5.2.1. Etkileşim Vekili (Engagement Proxy) Hesaplanması', bold=True, indent=False)
para('Yalnızca takipçi sayısına dayalı değerlendirmeler yanıltıcı olabilmektedir. Bu nedenle her içerik üreticisi için iki bileşenli bir Etkileşim Vekili skoru türetilmiştir. Bileşen 1 — Takipçi Başına Beğeni Oranı: like_per_follower = avg_video_likes / max(followers, 1). Bileşen 2 — Ortalama Video Beğeni Logaritması: log_avg_likes = ln(1 + avg_video_likes). Etkileşim Vekili (Normalize Edilmiş): Engagement_Proxy = 0.60 × MinMax(like_per_follower) + 0.40 × MinMax(log_avg_likes). %60 ağırlık, organik sadık kitleye sahip küçük üreticilerin adaletli biçimde temsil edilmesini sağlarken; %40 ağırlık, mutlak büyüklük avantajını değerlendirmeye katmaktadır.')

para('5.2.2. Tahmini Maliyet Modeli (Cost Estimation Model)', bold=True, indent=False)
para('TikTok işbirliği ücretleri kamuya açık değildir. Sektör benchmarkları ve uzman görüşlerine dayanan kademeli (tier-based) bir maliyet modeli geliştirilmiştir. Temel maliyet formülü: Maliyet = (Taban Ücret + Takipçi/1000 × Artış Oranı) × Kategori Çarpanı.')

tbl_caption('Tablo 5.1. Takipçi Bazlı Kademeli Maliyet Modeli')
t1 = doc.add_table(rows=6, cols=4); t1.style='Table Grid'; t1.alignment=WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Kademe (Tier)', 'Takipçi Aralığı', 'Taban Ücret (TL)', 'Artış Oranı (TL / 1K)']):
    cell_style(t1.rows[0].cells[i], h, bold=True, size=10, bg='D9D9D9')
for r, row in enumerate([('Nano','0 – 10.000','1.500','150'),('Micro','10.000 – 100.000','8.000','80'),
                          ('Mid','100.000 – 500.000','35.000','35'),('Macro','500.000 – 1.000.000','90.000','20'),
                          ('Mega','1.000.000+','200.000','8')]):
    for c, v in enumerate(row): cell_style(t1.rows[r+1].cells[c], v, size=10)
space()

tbl_caption('Tablo 5.2. Kategori Maliyet Çarpanları')
cats = [('Beauty & Personal Care (Güzellik ve Bakım)','1.15'),('Fashion & Style (Moda ve Stil)','1.10'),
        ('Technology & Digital (Teknoloji)','1.05'),('Travel & Lifestyle (Seyahat)','1.05'),
        ('Fitness & Health (Spor ve Sağlık)','1.00'),('Music & Performance (Müzik)','1.00'),
        ('Mixed / Unclear','1.00'),('Food & Cooking (Yemek)','0.95'),
        ('Comedy & Entertainment (Komedi)','0.90'),('Gaming (Oyun)','0.90'),
        ('Education & Informative (Eğitim)','0.85')]
t2 = doc.add_table(rows=len(cats)+1, cols=2); t2.style='Table Grid'; t2.alignment=WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Kategori','Maliyet Çarpanı']): cell_style(t2.rows[0].cells[i], h, bold=True, size=10, bg='D9D9D9')
for r, (c, m) in enumerate(cats):
    cell_style(t2.rows[r+1].cells[0], c, size=10); cell_style(t2.rows[r+1].cells[1], m, size=10)
space()

para('5.2.3. Çok Kriterli Karar Verme (MCDM) Skoru', bold=True, indent=False)
para('Birden fazla kriterin bütünleşik biçimde değerlendirilmesi için ağırlıklı doğrusal birleştirme yöntemi kullanılmıştır. Her kriter önce Min-Max normalizasyonuyla [0, 1] aralığına çekilmiş; ardından belirlenen ağırlıklarla ağırlıklı toplamı alınmıştır.')
tbl_caption('Tablo 5.3. MCDM Kriter Ağırlık Tablosu')
t3 = doc.add_table(rows=6, cols=3); t3.style='Table Grid'; t3.alignment=WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Kriter','Açıklama','Ağırlık']): cell_style(t3.rows[0].cells[i], h, bold=True, size=10, bg='D9D9D9')
for r, row in enumerate([('Etkileşim Skoru','Engagement Proxy','%35'),('Takipçi Büyüklüğü','log(1 + followers)','%20'),
                          ('Takipçi Başına Beğeni','like_per_follower','%20'),
                          ('Maliyet Etkinliği','log(avg_likes) / tahmini maliyet','%15'),
                          ('Video Hacmi','Video üretim sayısı','%10')]):
    for c, v in enumerate(row): cell_style(t3.rows[r+1].cells[c], v, size=10)
space()
para('Formül: MCDM_Score = 0.35 × c1 + 0.20 × c2 + 0.20 × c3 + 0.15 × c4 + 0.10 × c5. Etkileşim kalitesi (%55 birleşik ağırlıkla) en güçlü belirleyici unsur iken takipçi sayısının yalnızca büyüklüğü tek başına yeterli kabul edilmemektedir. Maliyet etkinliği (%15) bütçe kullanımını optimize ederken, video hacmi (%10) aktif içerik üretimini ödüllendirmektedir.')

heading('5.3. Optimizasyon Modeli: 0-1 Tam Sayılı Doğrusal Programlama', 2)
para('MCDM skorları ve maliyet tahminleri hesaplandıktan sonra, influencer seçim problemi 0-1 Tam Sayılı Doğrusal Programlama (ILP) modeli olarak kurulmuştur. ILP, karar değişkenlerinin tamsayı (0 veya 1) değerler almasını zorunlu kılan doğrusal optimizasyon modelidir. Influencer seçimi doğası gereği ikili bir karar yapısına sahip olduğundan bu yöntem probleme doğrudan uyum sağlamaktadır. Kim ve Cho [10] ile López-Dawn ve Giovanidis [11], benzer yapıdaki influencer seçim problemlerini matematiksel programlama yöntemiyle çözmenin kampanya etkinliğini anlamlı biçimde artırdığını göstermiştir.')
para('Karar Değişkeni: xᵢ ∈ {0, 1} — xᵢ = 1 → i. içerik üreticisi portföye dahil edilir; xᵢ = 0 → seçilmez.')
para('Amaç Fonksiyonu (Maksimizasyon): Z = Σ MCDM_Scoreᵢ × xᵢ → Seçilen influencer\'ların kümülatif MCDM skoru maksimize edilmektedir.')
para('Kısıt 1 — Bütçe Kısıtı: Σ costᵢ × xᵢ ≤ B → Seçilen influencer\'ların toplam tahmini maliyeti, kampanya bütçesini aşamaz.')
para('Kısıt 2 — Minimum Çeşitlilik: Σ xᵢ ≥ n_min → Risk dağıtımı için en az n_min influencer seçimi zorunludur.')
para('Kısıt 3 — Maksimum Portföy Büyüklüğü (İsteğe Bağlı): Σ xᵢ ≤ n_max → Yönetilebilirlik açısından üst sınır belirlenebilir.')
para('Kısıt 4 — Tier Bazlı Üst Sınır: Belirli kademedeki influencer sayısı sınırlandırılarak risk dengesi yönetilir.')
para('Kısıt 5 — Zorunlu Kategori Temsili: Kullanıcı tarafından seçilen kategorilerde en az bir influencer bulunması sağlanır.')
para('Kısıt 6 — Kara Liste (Blacklist): xₖ = 0, ∀k ∈ K_black → Kara listedeki influencer\'lar hiçbir koşulda seçilmez; bu kısıt markanın itibar yönetimi kararlarını matematiksel garanti altına alır.')
para('Model, Gurobi Solver (gurobipy Python kütüphanesi) ile çözülmektedir. Gurobi, Dal-ve-sınır (Branch-and-Bound) algoritması temelinde çalışan, akademik ve ticari kullanımda geniş alana sahip matematiksel programlama çözücüsüdür.')

tbl_caption('Tablo 5.4. Gurobi Çözücü Parametre Tablosu')
t4 = doc.add_table(rows=6, cols=3); t4.style='Table Grid'; t4.alignment=WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Parametre','Değer','Açıklama']): cell_style(t4.rows[0].cells[i], h, bold=True, size=10, bg='D9D9D9')
for r, row in enumerate([('TimeLimit','60 saniye','Maksimum çözüm süresi'),
                          ('MIPGap','0.001','%0.1 optimallik toleransı'),
                          ('Threads','4','Paralel iş parçacığı sayısı'),
                          ('MIPFocus','1','İyi çözüm bulmaya odaklan'),
                          ('Presolve','2','Agresif ön-çözme')]):
    for c, v in enumerate(row): cell_style(t4.rows[r+1].cells[c], v, size=10)
space()
para('Model çıktıları incelendiğinde, farklı bütçe seviyelerinde seçilen içerik üreticilerinin dağılımının değiştiği gözlemlenmiştir. Özellikle düşük bütçelerde daha yüksek etkileşim oranına sahip Micro influencer\'ların tercih edildiği, bütçe arttıkça daha geniş erişim sağlayan influencer\'ların da seçime dahil edildiği belirlenmiştir.')

heading('5.4. Streamlit Tabanlı Karar Destek Arayüzü', 2)
para('Yalnızca konsol çıktılarından ibaret akademik projeler ticari hayatta hızla geride kalmaktadır. Bu açığı kapatmak için Streamlit framework\'ü kullanılarak dinamik bir karar destek web uygulaması (streamlit_app.py) oluşturulmuştur.')
para('Bütçe Slaydırı aracılığıyla kampanya bütçesi 1.000 TL ile 10.000.000 TL arasında dinamik olarak belirlenebilir. Kategori Seçimi ile zorunlu içerik kategorileri çoklu seçim kutucuklarıyla tanımlanabilir. Kara Liste Yönetimi sayesinde istenmeyen influencer\'lar listeden seçilerek otomatik olarak model dışında tutulur; profil görselleri de görüntülenir. Optimizasyon Tetikleyicisi ile tek tuşa basılarak Gurobi modeli arka planda çalıştırılır. Sonuç Görüntüleme ekranında seçilen influencer sayısı, kullanılan bütçe, kalan bütçe ve toplam MCDM skoru metrik kartlarda gösterilir; seçilen influencer\'lar ızgara düzeninde fotoğrafları, kategorileri, takipçi sayıları, tahmini maliyetleri ve MCDM skorlarıyla listelenir.')
para('Bu arayüz tasarımı, teknik bilgiden bağımsız olarak bir pazarlama yöneticisinin modeli doğrudan kullanabilmesini sağlamaktadır.')

pb()

# ════════════════════════════════════════════════════════════════════
# 6. SONUÇ
# ════════════════════════════════════════════════════════════════════
heading('6. SONUÇ VE ÖNERİLER', 1)
para('Bu tez çalışmasının ENM458 kapsamındaki ilk altı haftasında gerçekleştirilen çalışmalar ve ulaşılan sonuçlar şu şekilde özetlenebilir.')
para('Birinci dönemde (ENM457) Selenium tabanlı web kazıma yöntemiyle oluşturulan TikTok veri seti, ikinci dönemde (ENM458) analitik dönüşüm süreçlerinden geçirilerek optimizasyon modeline hazır hale getirilmiştir. Etkileşim vekili hesaplamaları, takipçi sayısının tek başına performansı açıklamadaki yetersizliğini bir kez daha doğrulamıştır; yüksek takipçili bazı hesapların, düşük takipçili ama etkileşim odaklı hesapların gerisinde MCDM skoru aldığı gözlemlenmiştir.')
para('Gurobi ILP modeli, farklı bütçe senaryolarında (50K, 150K, 300K, 600K TL) tutarlı ve matematiksel olarak optimal portföyler üretmektedir. Model, geleneksel sezgisel yaklaşımların aksine portföy bütçesini katmanlar arasında dengeleyerek hem maliyet etkinliğini hem de kategori çeşitliliğini optimize etmektedir. Streamlit arayüzü, modeli teknik bilgisi bulunmayan son kullanıcılara başarıyla sunmaktadır.')
para('Sınırlılıklar: Maliyet modeli, gerçek piyasa fiyatlarının sektör benchmarklarına dayanan bir tahminini temsil etmekte; bireysel pazarlık koşulları yansıtılmamaktadır. Veri seti yalnızca belirli bir zaman dilimini kapsamakta olup influencer performansı zaman içinde dalgalanabilmektedir. Web kazıma yöntemi yalnızca herkese açık profillerle sınırlıdır. Birinci dönem raporunda "herhangi bir optimizasyon uygulanmamıştır" olarak belirtilen olgu, ENM458 kapsamında tamamıyla aşılmış; MCDM skorlama ve Gurobi ILP modeli hayata geçirilmiş ve Streamlit arayüzüne entegre edilmiştir.')
para('Sonraki Aşamada Planlanlar (Hafta 7–14): Farklı bütçe seviyeleri ve kriter ağırlık kombinasyonları için kapsamlı duyarlılık analizi (Sensitivity Analysis) gerçekleştirilecektir. Modelin farklı marka profilleri için örnek vaka çalışmaları hazırlanacaktır. Streamlit arayüzüne senaryo karşılaştırma ve sonuç dışa aktarma (Excel/PDF) özellikleri eklenecektir. Tüm çalışma ENM458 tez yazım kılavuzuna uygun biçimde formatlanarak nihai forma taşınacaktır.')
para('Sonuç olarak bu çalışma, TikTok platformu özelinde içerik üreticisi pazarlamasına yönelik veri temelli bir optimizasyon sistemi sunarak içerik üreticisi seçiminin daha bilinçli, ölçülebilir ve sistematik biçimde gerçekleştirilmesine katkı sağlamaktadır.')

pb()

# ════════════════════════════════════════════════════════════════════
# KAYNAKÇA
# ════════════════════════════════════════════════════════════════════
heading('KAYNAKÇA', 1)
refs = [
    '[1]\tJohnson, M., & Alvarez, L. (2021). The Rise of Digital Influencers: Shaping the Future of Marketing. Journal of Digital Communication, 14(3), 210–225.',
    '[2]\tKumar, R., & Lee, D. (2020). Influencer Marketing with Social Platforms. Social Media Research Journal, 8(4), 112–130.',
    '[3]\tZhang, P., & Han, Y. (2023). Social Media Influencer Marketing: Foundations, Trends and Research Directions. International Journal of Marketing Science, 19(1), 55–74.',
    '[4]\tAndersson, S., & Bergström, E. (2020). Instagram and Influencer Marketing: An Empirical Study of the Parameters Behind Success. Procedia Economics and Business, 7(2), 89–101.',
    '[5]\tLi, C., & Wang, H. (2022). Computational Studies in Influencer Marketing. Expert Systems with Applications, 193, 116–127.',
    '[6]\tAraujo, T., Neijens, P., & Vliegenthart, R. (2019). Discovering Effective Influencers. Computers in Human Behavior, 98, 10–20.',
    '[7]\tSaito, M., & Kobayashi, K. (2019). A SI Model for Social Media Influencer Maximization. IEEE Access, 7, 150876–150889.',
    '[8]\tTiukhova, L., Korovin, D., & Melnikov, P. (2022). Influencer Prediction with Dynamic Graph Neural Networks. Neural Networks, 154, 145–159.',
    '[9]\tInfluencer Marketing Hub. (2023). Influencer Marketing Benchmark Report 2023. Global Industry Review Series.',
    '[10]\tKim, S., & Cho, Y. (2023). Optimal Influencer Marketing Campaign under Budget Constraints. Journal of Business Analytics, 12(3), 180–195.',
    '[11]\tLópez-Dawn, A., & Giovanidis, A. (2021). Budgeted Portfolio Optimization Model for Social Media Influencer Selection. Journal of Applied Optimization, 18(4), 300–315.',
    '[12]\tDe Veirman, M., Cauberghe, V., & Hudders, L. (2017). Marketing through Instagram influencers: The impact of number of followers and product divergence. International Journal of Advertising, 36(5), 798–828.',
    '[13]\tLou, C., & Yuan, S. (2019). Influencer marketing: How message value and credibility affect consumer trust. Journal of Interactive Advertising, 19(1), 58–73.',
    '[14]\tPhua, J., Jin, S. V., & Kim, J. (2020). The roles of celebrity endorsers\' credibility and attractiveness in influencer marketing. Computers in Human Behavior, 102, 310–321.',
    '[15]\tDwivedi, Y. K., et al. (2021). Setting the future of digital and social media marketing research. International Journal of Information Management, 59, 102168.',
]
for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = Pt(18)
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent  = Cm(1)
    r = p.add_run(ref); r.font.name='Times New Roman'; r.font.size=Pt(12)

pb()

# ════════════════════════════════════════════════════════════════════
# EKLER
# ════════════════════════════════════════════════════════════════════
heading('7. EKLER', 1)
tbl_caption('Tablo E.1. İş-Zaman Çizelgesi')
t5 = doc.add_table(rows=8, cols=4); t5.style='Table Grid'; t5.alignment=WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['No','İş Paketinin Adı','Zaman Aralığı','Takım Lideri']):
    cell_style(t5.rows[0].cells[i], h, bold=True, size=10, bg='D9D9D9')
jobs = [
    ('1','Veri Toplama Stratejisinin Belirlenmesi ve Ön Analiz','Ekim – Kasım 2025','Esra Zeyrek'),
    ('2','Kapsamlı Verinin Toplanması ve Veri Temizleme Süreci','Kasım – Aralık 2025','A. Raif Yıldırım'),
    ('3','Toplanan Verilerin İşlenmesi ve İstatistiksel Değerlendirme','Aralık 2025 – Ocak 2026','Oğuzhan Dikmen'),
    ('4','Engagement Proxy, MCDM Skoru ve Maliyet Modeli Geliştirme','Ocak – Şubat 2026','A. Raif Yıldırım'),
    ('5','ILP Modelinin Kurulması, Gurobi Entegrasyonu ve Streamlit Arayüzü','Şubat – Mart 2026','Esra Zeyrek'),
    ('6','Model Çözümü, Test Süreci ve Duyarlılık Analizi','Nisan – Mayıs 2026','Oğuzhan Dikmen'),
    ('7','Sonuçların Değerlendirilmesi ve Raporlama','Mayıs 2026','Esra Zeyrek'),
]
for r, row in enumerate(jobs):
    for c, v in enumerate(row): cell_style(t5.rows[r+1].cells[c], v, size=10)

space()
tbl_caption('Tablo E.2. Risk Yönetim Tablosu (B Planı)')
t6 = doc.add_table(rows=8, cols=3); t6.style='Table Grid'; t6.alignment=WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['No','En Büyük Riskler','Risk Yönetimi (B Planı)']):
    cell_style(t6.rows[0].cells[i], h, bold=True, size=10, bg='D9D9D9')
risks = [
    ('1','TikTok\'un veri erişimini kısıtlaması veya doğrulama mekanizmalarını artırması','Esnek scraper yapısı kullanılmaktadır. Kısıtlama artarsa veri çekme hızı azaltılacak, veri seti ek kaynaklarla zenginleştirilecektir.'),
    ('2','Toplanan verinin analiz için yetersiz kalması','Örneklem genişletilebilir biçimde tasarlanmıştır. Yetersizlik durumunda kullanıcı sayısı artırılacak veya normalize metrikler kullanılacaktır.'),
    ('3','API eksikliği nedeniyle veri tutarlılığının sağlanamaması','Veriler yalnızca herkese açık kaynaklardan toplanmış, temizleme ve çapraz doğrulama ile tutarlılık güvence altına alınmıştır.'),
    ('4','Web kazıma sürecinde teknik hatalar veya kesintiler yaşanması','Checkpoint mekanizmalarıyla veri kaybı önlenmiştir. Hata durumunda süreç kaldığı yerden devam eder.'),
    ('5','Optimizasyon modelinde uygulanabilir çözüm bulunamaması (infeasible)','Bütçe artırılacak, kategori kısıtları esnetilecek veya kara liste daraltılacaktır. Arayüz kullanıcıya öneride bulunmaktadır.'),
    ('6','Gurobi lisans sorunları yaşanması','PuLP + CBC (açık kaynak) çözücüsüne geçiş B planı olarak hazır tutulmaktadır; model bu geçişe uyumlu kodlanmıştır.'),
    ('7','Zaman planına uyulamaması','İş paketleri paralel planlanmıştır. Gecikme durumunda kapsam daraltılarak temel model korunacak, gelişmiş özellikler sonraki versiyona aktarılacaktır.'),
]
for r, row in enumerate(risks):
    for c, v in enumerate(row): cell_style(t6.rows[r+1].cells[c], v, size=10)

path = 'ENM458_Ara_Rapor_tiktok.docx'
doc.save(path)
print(f'✅ Kaydedildi: {path}')
